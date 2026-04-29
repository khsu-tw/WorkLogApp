"""
Work Log Journal  — Cross-Platform Web App
==========================================
Runs on Windows, macOS, iOS (Safari PWA), Android

Setup:
    pip install flask pocketbase python-docx Pillow

Cloud DB (PocketBase — free tier):
    1. Create account at https://pocketbase.com
    2. New project → SQL Editor → run schema.sql
    3. Copy Project URL + anon key into .env or config below

Run:
    python app.py
    Then open http://localhost:5000 on any device on the same network.
    On iOS: Safari → Share → "Add to Home Screen" for PWA install.
"""

import sys
import os, re, io, json, base64, hashlib, datetime, tempfile
from pathlib import Path

# ── Read version from VERSION file (single source of truth) ──────────────────
def _read_version() -> str:
    try:
        return (Path(__file__).parent / "VERSION").read_text(encoding="utf-8").strip()
    except (FileNotFoundError, OSError):
        return "0.0.0"

APP_VERSION = _read_version()

# ── Load .env file (must happen before os.environ.get) ──────────────────────
try:
    from dotenv import load_dotenv
    _env_path = Path(__file__).parent / ".env"
    load_dotenv(dotenv_path=_env_path, override=True)
except ImportError:
    pass  # python-dotenv not installed — use os environment only

from flask import Flask, request, jsonify, render_template_string, send_file, make_response

# ── Optional cloud DB (PocketBase) ──────────────────────────────────────────
try:
    from pocketbase import PocketBase
    HAS_POCKETBASE = True
except ImportError:
    HAS_POCKETBASE = False

# ── Optional cloud DB (PostgreSQL) - v1.0.1 ─────────────────────────────────
try:
    import psycopg2
    from psycopg2.extras import RealDictCursor
    HAS_POSTGRES = True
except ImportError:
    HAS_POSTGRES = False

import requests  # For PocketBase health check

# ── Optional Word export ─────────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ── Optional Excel export ─────────────────────────────────────────────────────
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

# ── Optional PDF export ──────────────────────────────────────────────────────
try:
    from fpdf import FPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

# ── Optional image support ───────────────────────────────────────────────────
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

import sqlite3

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB Flask default
app.secret_key = os.environ.get("SECRET_KEY", "worklog-dev-key-change-in-prod")

# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURATION  — edit these or set as environment variables
# ─────────────────────────────────────────────────────────────────────────────
POCKETBASE_URL = os.environ.get("POCKETBASE_URL", "http://127.0.0.1:8090")
POCKETBASE_TOKEN = os.environ.get("POCKETBASE_TOKEN", "")
# PostgreSQL configuration (v1.0.1 - alternative to PocketBase)
POSTGRES_URL = os.environ.get("POSTGRES_URL", "")  # Format: postgresql://user:pass@host:port/db
TABLE = "worklog"

# Resolve DB path relative to the executable (works both packaged and in dev)
def _resolve_db_path() -> str:
    custom = os.environ.get("SQLITE_PATH", "")
    if custom:
        return custom
    if getattr(sys, 'frozen', False):
        # PyInstaller bundle: store next to the .exe
        base = Path(sys.executable).parent
    else:
        base = Path(__file__).parent
    return str(base / "WorkLog.db")

SQLITE_PATH = _resolve_db_path()

STATUS_OPTIONS   = ["Not Started","Early Engagement","WIP","Production",
                    "Cancelled","No More Activity","Task Done"]
CATEGORY_OPTIONS = ["General","Visit","Training","Tech Support","Design Review",
                    "Debug","Documentation","Design-in","Evaluation","Follow-up",
                    "Design-win","Others"]

# ─────────────────────────────────────────────────────────────────────────────
# DATABASE — OFFLINE-FIRST HYBRID
# ─────────────────────────────────────────────────────────────────────────────
"""
Architecture:
  LocalDB  — always used for reads/writes (SQLite, always fast, works offline)
  CloudDB  — PocketBase, used only when online
  SyncEngine — background thread; pushes pending_sync rows to cloud,
               detects conflicts by comparing last_update timestamps,
               pulls new/updated cloud rows not yet in local cache.

Every write to LocalDB stamps the row with:
  sync_status: 'pending' | 'synced' | 'conflict'
  local_updated_at: ISO timestamp of local write (for conflict detection)
"""

import threading as _threading

_sync_lock = _threading.Lock()


class LocalDB:
    """
    SQLite — always the primary read/write store.
    Adds sync_status + local_updated_at columns for offline tracking.
    """
    def __init__(self, path=None):
        self.path = path or SQLITE_PATH
        self._init()

    def _conn(self):
        c = sqlite3.connect(self.path, check_same_thread=False)
        c.row_factory = sqlite3.Row
        c.execute("PRAGMA journal_mode=WAL")
        return c

    def _init(self):
        c = self._conn()
        try:
            c.execute("""CREATE TABLE IF NOT EXISTS worklog (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                week TEXT, due_date TEXT, customer TEXT, project_name TEXT,
                sso_modeln TEXT, ear TEXT, application TEXT, bu TEXT,
                task_summary TEXT, mchp_device TEXT, project_schedule TEXT,
                todo_content TEXT, todo_due_date TEXT,
                status TEXT, category TEXT,
                worklogs TEXT, create_date TEXT, last_update TEXT,
                archive TEXT DEFAULT 'No', record_hash TEXT,
                cloud_id TEXT,
                sync_status TEXT DEFAULT 'pending',
                local_updated_at TEXT
            )""")
            c.execute("""CREATE TABLE IF NOT EXISTS sync_conflicts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                local_id INTEGER,
                cloud_snapshot TEXT,
                detected_at TEXT,
                conflict_type TEXT DEFAULT 'normal'
            )""")
            cur = c.execute("PRAGMA table_info(worklog)")
            cols = {r[1] for r in cur.fetchall()}
            for col, dtype in [
                ("due_date","TEXT"),("record_hash","TEXT"),("category","TEXT"),
                ("mchp_device","TEXT"),("cloud_id","TEXT"),
                ("sync_status","TEXT"),("local_updated_at","TEXT"),
                ("sso_modeln","TEXT"),("ear","TEXT"),("application","TEXT"),("bu","TEXT"),
                ("project_schedule","TEXT"),
                ("todo_content","TEXT"),("todo_due_date","TEXT"),
            ]:
                if col not in cols:
                    c.execute(f"ALTER TABLE worklog ADD COLUMN {col} {dtype}")
            # Migrate sync_conflicts table
            cur2 = c.execute("PRAGMA table_info(sync_conflicts)")
            conflict_cols = {r[1] for r in cur2.fetchall()}
            if "conflict_type" not in conflict_cols:
                c.execute("ALTER TABLE sync_conflicts ADD COLUMN conflict_type TEXT DEFAULT 'normal'")
            c.commit()
        finally:
            c.close()

    def fetch_all(self, filters=None):
        sql = "SELECT * FROM worklog WHERE IFNULL(sync_status, 'pending') != 'deleted'"
        params = []
        if filters:
            clauses = []
            for col, val in filters.items():
                if val:
                    if col in ("customer","project_name"):
                        clauses.append(f"{col} LIKE ?"); params.append(f"%{val}%")
                    elif col == "archive":
                        clauses.append(f"IFNULL({col}, 'No')=?"); params.append(val)
                    else:
                        clauses.append(f"{col}=?"); params.append(val)
            if clauses:
                sql += " AND " + " AND ".join(clauses)
        sql += " ORDER BY id DESC"
        c = self._conn()
        try:
            return [dict(r) for r in c.execute(sql, params).fetchall()]
        finally:
            c.close()

    def _now_iso(self):
        return datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.%f")

    def insert(self, data):
        data = dict(data)
        data.setdefault("sync_status", "pending")
        data["local_updated_at"] = self._now_iso()
        cols = ", ".join(data.keys())
        phs  = ", ".join(["?"]*len(data))
        c = self._conn()
        try:
            cur = c.execute(f"INSERT INTO worklog ({cols}) VALUES ({phs})",
                            list(data.values()))
            c.commit()
            result = {**data, "id": cur.lastrowid}
            # Trigger immediate sync on data change
            _trigger_sync()
            return result
        except Exception as e:
            c.rollback(); raise e
        finally:
            c.close()

    def update(self, row_id, data):
        data = dict(data)
        data["sync_status"]     = "pending"
        data["local_updated_at"] = self._now_iso()
        sets = ", ".join([f"{k}=?" for k in data])
        c = self._conn()
        try:
            c.execute(f"UPDATE worklog SET {sets} WHERE id=?",
                      list(data.values()) + [row_id])
            c.commit()
            # Trigger immediate sync on data change
            _trigger_sync()
        except Exception as e:
            c.rollback(); raise e
        finally:
            c.close()
        return {**data, "id": row_id}

    def delete(self, row_id):
        """Soft-delete: mark as 'deleted' so sync can propagate to cloud."""
        c = self._conn()
        try:
            c.execute("UPDATE worklog SET sync_status='deleted' WHERE id=?", (row_id,))
            c.commit()
            # Trigger immediate sync on data change
            _trigger_sync()
        finally:
            c.close()

    def find_by_hash(self, h):
        c = self._conn()
        try:
            rows = c.execute(
                "SELECT * FROM worklog WHERE record_hash=? AND sync_status!='deleted'",
                (h,)).fetchall()
            return dict(rows[0]) if rows else None
        finally:
            c.close()

    def find_by_cloud_id(self, cloud_id):
        c = self._conn()
        try:
            rows = c.execute("SELECT * FROM worklog WHERE cloud_id=?",
                             (str(cloud_id),)).fetchall()
            return dict(rows[0]) if rows else None
        finally:
            c.close()

    def mark_synced(self, local_id, cloud_id):
        c = self._conn()
        try:
            c.execute("UPDATE worklog SET sync_status='synced', cloud_id=? WHERE id=?",
                      (str(cloud_id), local_id))
            c.commit()
        finally:
            c.close()

    def mark_conflict(self, local_id, cloud_snapshot: dict, conflict_type: str = "normal"):
        """Store the conflicting cloud version as JSON in a side table.
        conflict_type: 'normal' (timestamp-based) or 'unresolvable' (requires user decision)

        v1.0.5: Enhanced with field-by-field diff detection.
        """
        c = self._conn()
        try:
            # Get current local version for field comparison
            local_row = c.execute("SELECT * FROM worklog WHERE id=?", (local_id,)).fetchone()
            if local_row:
                local_dict = dict(local_row)
                # Calculate field differences
                diff_fields = self._calculate_field_diff(local_dict, cloud_snapshot)
                cloud_snapshot["_diff_fields"] = diff_fields  # Store diff metadata

            c.execute(
                "INSERT INTO sync_conflicts (local_id, cloud_snapshot, detected_at, conflict_type) VALUES (?,?,?,?)",
                (local_id, json.dumps(cloud_snapshot),
                 datetime.datetime.now(datetime.timezone.utc).isoformat(), conflict_type))
            c.execute("UPDATE worklog SET sync_status='conflict' WHERE id=?", (local_id,))
            c.commit()
        finally:
            c.close()

    def _calculate_field_diff(self, local: dict, cloud: dict) -> list:
        """Compare local and cloud versions field-by-field.
        Returns list of field names that differ (v1.0.5).
        """
        ignore_fields = {"id", "cloud_id", "sync_status", "local_updated_at", "created", "updated"}
        diff_fields = []

        # Map cloud "archieved" to local "archive" for comparison
        cloud_normalized = dict(cloud)
        if "archieved" in cloud_normalized:
            cloud_normalized["archive"] = cloud_normalized.pop("archieved")

        all_fields = set(local.keys()) | set(cloud_normalized.keys())
        for field in all_fields:
            if field in ignore_fields:
                continue
            local_val = str(local.get(field, "")).strip()
            cloud_val = str(cloud_normalized.get(field, "")).strip()
            if local_val != cloud_val:
                diff_fields.append(field)

        return diff_fields

    def get_conflicts(self):
        c = self._conn()
        try:
            rows = c.execute("""
                SELECT sc.id as conflict_id, sc.local_id, sc.cloud_snapshot,
                       sc.detected_at, sc.conflict_type, w.*
                FROM sync_conflicts sc
                JOIN worklog w ON w.id = sc.local_id
                ORDER BY sc.detected_at DESC
            """).fetchall()
            result = []
            for r in rows:
                d = dict(r)
                d["cloud_snapshot"] = json.loads(d["cloud_snapshot"] or "{}")
                result.append(d)
            return result
        finally:
            c.close()

    def resolve_conflict(self, conflict_id: int, keep: str):
        """keep: 'local', 'cloud', or 'backup' (keep both versions)"""
        c = self._conn()
        try:
            row = dict(c.execute(
                "SELECT * FROM sync_conflicts WHERE id=?", (conflict_id,)).fetchone())
            local_id      = row["local_id"]
            cloud_snapshot = json.loads(row["cloud_snapshot"])

            if keep == "cloud":
                # Overwrite local with cloud version, mark synced
                snap = dict(cloud_snapshot)
                snap.pop("id", None)
                snap["sync_status"]     = "synced"
                snap["local_updated_at"] = datetime.datetime.now(datetime.timezone.utc).isoformat()
                snap["cloud_id"]        = str(cloud_snapshot.get("id",""))
                sets = ", ".join([f"{k}=?" for k in snap])
                c.execute(f"UPDATE worklog SET {sets} WHERE id=?",
                          list(snap.values()) + [local_id])
            elif keep == "backup":
                # Keep local as-is, and insert cloud version as a new backup record
                snap = dict(cloud_snapshot)
                snap.pop("id", None)
                snap.pop("created_at", None)
                snap["sync_status"]     = "synced"
                snap["local_updated_at"] = datetime.datetime.now(datetime.timezone.utc).isoformat()
                snap["cloud_id"]        = str(cloud_snapshot.get("id",""))
                # Add "(Backup)" suffix to identify backup copies
                if snap.get("customer"):
                    snap["customer"] = snap["customer"] + " (Backup)"
                cols = ", ".join(snap.keys())
                phs  = ", ".join(["?"]*len(snap))
                c.execute(f"INSERT INTO worklog ({cols}) VALUES ({phs})", list(snap.values()))
                # Mark original local as pending
                c.execute("UPDATE worklog SET sync_status='pending' WHERE id=?", (local_id,))
            else:
                # Keep local, mark pending so it gets pushed to cloud
                c.execute(
                    "UPDATE worklog SET sync_status='pending' WHERE id=?", (local_id,))

            c.execute("DELETE FROM sync_conflicts WHERE id=?", (conflict_id,))
            c.commit()
        finally:
            c.close()

    def pending_rows(self):
        c = self._conn()
        try:
            return [dict(r) for r in c.execute(
                "SELECT * FROM worklog WHERE sync_status IN ('pending','deleted')"
            ).fetchall()]
        finally:
            c.close()


# ── Single shared LocalDB instance ────────────────────────────────────────────
_local_db = None

def _get_local_db() -> LocalDB:
    global _local_db, SQLITE_PATH
    # Re-resolve path in case env was updated by launcher after module load
    SQLITE_PATH = _resolve_db_path()
    if _local_db is None or _local_db.path != SQLITE_PATH:
        _local_db = LocalDB(path=SQLITE_PATH)
    return _local_db


def _has_local_data() -> bool:
    """True if the local DB already has at least one non-deleted record."""
    try:
        return len(_get_local_db().fetch_all()) > 0
    except Exception:
        return False


# ── Cloud connectivity check ──────────────────────────────────────────────────
def _is_online() -> bool:
    """True when cloud database (PocketBase or PostgreSQL) is reachable."""
    # Check PocketBase first
    if _pocketbase_configured():
        url = os.environ.get("POCKETBASE_URL", POCKETBASE_URL)
        try:
            resp = requests.get(f"{url.rstrip('/')}/api/health", timeout=3)
            if resp.status_code == 200:
                return True
        except Exception:
            pass
    if _postgres_configured():
        return _postgres_online()
    return False


def _cloud_client():
    """Create PocketBase client with optional token authentication."""
    url = os.environ.get("POCKETBASE_URL", POCKETBASE_URL)
    pb = PocketBase(url)
    token = os.environ.get("POCKETBASE_TOKEN", POCKETBASE_TOKEN)
    if token:
        pb.auth_store.save(token)
    return pb


def _pocketbase_configured() -> bool:
    """Check if PocketBase is configured."""
    return bool(
        HAS_POCKETBASE
        and os.environ.get("POCKETBASE_URL", POCKETBASE_URL)
    )


# ── PostgreSQL support ────────────────────────────────────────────────────────
def _postgres_configured() -> bool:
    """Check if PostgreSQL is configured."""
    return bool(HAS_POSTGRES and os.environ.get("POSTGRES_URL", POSTGRES_URL))


def _postgres_client():
    """Create PostgreSQL connection."""
    url = os.environ.get("POSTGRES_URL", POSTGRES_URL)
    if not url:
        return None
    try:
        conn = psycopg2.connect(url, cursor_factory=RealDictCursor, connect_timeout=5)
        return conn
    except Exception:
        return None


def _postgres_online() -> bool:
    """Check if PostgreSQL is reachable."""
    if not _postgres_configured():
        return False
    try:
        conn = _postgres_client()
        if conn:
            conn.close()
            return True
    except Exception:
        pass
    return False


# ── Sync engine ───────────────────────────────────────────────────────────────
class SyncEngine:
    """
    Background worker:
      • Every 60s: checks connectivity, pushes pending rows, pulls remote changes.
      • Smart sync: skips sync cycle when no pending changes (saves bandwidth).
      • On data change: immediate push to cloud.
      • Conflict = local pending AND cloud version has a newer last_update.
        → stores conflict, flags row; user resolves via /api/sync/conflicts.
    """
    INTERVAL = 60   # seconds between sync attempts (1 minute, v1.0.5)

    def __init__(self):
        self._stop   = _threading.Event()
        self._thread = _threading.Thread(target=self._loop, daemon=True)
        self._status = {"online": False, "last_sync": None,
                        "pending": 0, "conflicts": 0}

    def start(self):
        self._thread.start()
        # On first start: if we have PocketBase configured and no local data yet,
        # do an immediate pull so the user sees cloud data right away.
        _threading.Thread(target=self._initial_pull, daemon=True).start()

    def _create_cloud_backup(self):
        """Create a JSON backup of all cloud data in the local directory."""
        try:
            pb = _cloud_client()
            if not pb:
                return
            records = pb.collection(TABLE).get_full_list()
            cloud_rows = [self._pb_record_to_dict(r) for r in records]
            if cloud_rows:
                backup_dir = os.path.dirname(SQLITE_PATH)
                backup_file = os.path.join(backup_dir, f"pocketbase_backup_{datetime.date.today().strftime('%Y%m%d')}.json")
                with open(backup_file, "w", encoding="utf-8") as f:
                    json.dump(cloud_rows, f, ensure_ascii=False, indent=2, default=str)
                print(f"  Cloud backup created: {backup_file}")
        except Exception as e:
            print(f"  Cloud backup failed: {e}")

    def _initial_pull(self):
        """Pull all cloud data into local on first launch (empty local DB)."""
        import time as _time
        # Give Flask a moment to finish starting
        _time.sleep(2)
        if _is_online():
            # Always create a backup of cloud data when online (PocketBase only)
            if not _postgres_configured():
                self._create_cloud_backup()
            if not _has_local_data():
                report = {"pushed": 0, "pulled": 0, "conflicts": 0, "errors": []}
                with _sync_lock:
                    try:
                        if _postgres_configured():
                            self._pull_postgres(report)
                        else:
                            self._pull(report)
                        if report["pulled"]:
                            self._status["last_sync"] = datetime.datetime.now().strftime(
                                "%Y-%m-%d %H:%M:%S")
                            self._status["online"] = True
                    except Exception as e:
                        report["errors"].append(str(e))

    def stop(self):
        self._stop.set()

    def status(self) -> dict:
        db = _get_local_db()
        pending   = len([r for r in db.pending_rows() if r["sync_status"]=="pending"])
        conflicts = len(db.get_conflicts())
        self._status.update({"pending": pending, "conflicts": conflicts})
        return dict(self._status)

    def _loop(self):
        while not self._stop.wait(self.INTERVAL):
            # Smart sync: skip if no pending changes (v1.0.5 bandwidth optimization)
            db = _get_local_db()
            pending_count = len(db.pending_rows())
            if pending_count > 0:
                self.sync_now()
            else:
                # Still update online status for UI
                if _is_online():
                    self._status["online"] = True

    def sync_now(self) -> dict:
        if not _is_online():
            self._status["online"] = False
            return {"ok": False, "reason": "offline"}

        self._status["online"] = True
        report = {"pushed": 0, "pulled": 0, "conflicts": 0, "errors": []}

        with _sync_lock:
            try:
                if _postgres_configured():
                    self._push_postgres(report)
                    self._pull_postgres(report)
                else:
                    self._push(report)
                    self._pull(report)
                self._status["last_sync"] = datetime.datetime.now().strftime(
                    "%Y-%m-%d %H:%M:%S")
                self._status["pending"]   = 0
            except Exception as e:
                report["errors"].append(str(e))

        return report

    # ── Push pending local rows to PocketBase ─────────────────────────────────
    def _push(self, report: dict):
        db = _get_local_db()
        pb = _cloud_client()
        pending = db.pending_rows()

        for row in pending:
            local_id = row["id"]
            cloud_id = row.get("cloud_id")

            # Soft-deleted rows
            if row["sync_status"] == "deleted":
                if cloud_id:
                    try:
                        pb.collection(TABLE).delete(cloud_id)
                    except Exception:
                        pass
                c = db._conn()
                c.execute("DELETE FROM worklog WHERE id=?", (local_id,))
                c.commit(); c.close()
                report["pushed"] += 1
                continue

            # Prepare payload — strip local-only columns
            payload = {k: v for k, v in row.items()
                       if k not in ("id","cloud_id","sync_status","local_updated_at")}

            # Map local "archive" → PocketBase "archieved"
            if "archive" in payload:
                payload["archieved"] = payload.pop("archive")

            if cloud_id:
                # Check for conflict: compare timestamps (use local_updated_at for precision)
                try:
                    cloud_row = pb.collection(TABLE).get_one(cloud_id)
                    cloud_lu = getattr(cloud_row, 'last_update', '') or ''
                    local_lu = (row.get("last_update") or "")
                    local_updated = (row.get("local_updated_at") or "")

                    # Enhanced conflict resolution based on timestamps
                    if cloud_lu and local_lu:
                        if cloud_lu > local_lu:
                            # Cloud is newer - flag for user
                            cloud_data = {k: getattr(cloud_row, k, None) for k in dir(cloud_row) if not k.startswith('_')}
                            db.mark_conflict(local_id, cloud_data)
                            report["conflicts"] += 1
                            continue
                        elif cloud_lu == local_lu and local_updated:
                            # Same last_update but local has been modified - local wins
                            pass  # Continue to push
                    elif not cloud_lu and not local_lu:
                        # Both timestamps missing - flag as unresolvable conflict
                        cloud_data = {k: getattr(cloud_row, k, None) for k in dir(cloud_row) if not k.startswith('_')}
                        db.mark_conflict(local_id, cloud_data, conflict_type="unresolvable")
                        report["conflicts"] += 1
                        continue
                except Exception:
                    pass  # Record not found in cloud, will be created

                try:
                    pb.collection(TABLE).update(cloud_id, payload)
                    db.mark_synced(local_id, cloud_id)
                    report["pushed"] += 1
                except Exception as e:
                    report["errors"].append(f"update {local_id}: {e}")
            else:
                try:
                    result = pb.collection(TABLE).create(payload)
                    if result:
                        new_cloud_id = result.id
                        db.mark_synced(local_id, new_cloud_id)
                        report["pushed"] += 1
                except Exception as e:
                    report["errors"].append(f"insert {local_id}: {e}")

    # ── Pull cloud rows into local cache ──────────────────────────────────────
    def _normalize_cloud_row(self, row: dict) -> dict:
        """Map PocketBase record to local schema."""
        # Map old 'inquiries' to new 'mchp_device' if exists
        if "inquiries" in row and "mchp_device" not in row:
            row["mchp_device"] = row.pop("inquiries")
        # Ensure new columns have defaults if missing from cloud
        row.setdefault("ear", "")
        row.setdefault("bu", "")
        row.setdefault("mchp_device", "")
        row.setdefault("sso_modeln", "")
        row.setdefault("application", "")
        row.setdefault("project_schedule", "")
        row.setdefault("todo_content", "")
        row.setdefault("todo_due_date", "")
        row.setdefault("archive", "No")
        return row

    def _pb_record_to_dict(self, record) -> dict:
        """Convert PocketBase record object to dictionary."""
        result = {}
        for field in ["id", "week", "due_date", "customer", "project_name", "sso_modeln",
                      "ear", "application", "bu", "task_summary", "mchp_device",
                      "project_schedule", "todo_content", "todo_due_date", "status", "category", "worklogs", "create_date",
                      "last_update", "record_hash", "created", "updated"]:
            result[field] = getattr(record, field, None)
        result["archive"] = getattr(record, "archieved", None) or "No"
        return result

    def _pull(self, report: dict):
        db = _get_local_db()
        pb = _cloud_client()

        # Fetch ALL cloud rows from PocketBase
        try:
            records = pb.collection(TABLE).get_full_list()
            cloud_rows = [self._pb_record_to_dict(r) for r in records]
        except Exception as e:
            report["errors"].append(f"pull fetch: {e}")
            return

        for crow in cloud_rows:
            crow = self._normalize_cloud_row(crow)
            cloud_id = str(crow["id"])
            local = db.find_by_cloud_id(cloud_id)

            if local is None:
                # New cloud row not yet in local
                payload = {k: v for k, v in crow.items()
                           if k not in ("created", "updated")}
                payload.pop("id", None)
                payload["cloud_id"] = cloud_id
                payload["sync_status"] = "synced"
                payload["local_updated_at"] = datetime.datetime.now(datetime.timezone.utc).isoformat()
                # Remove unknown columns
                for col in ("created", "updated", "inquiries"):
                    payload.pop(col, None)
                try:
                    db.insert(payload)
                    report["pulled"] += 1
                except Exception as e:
                    report["errors"].append(f"pull insert {cloud_id}: {e}")
            elif local["sync_status"] == "synced":
                cloud_lu = crow.get("last_update") or ""
                local_lu = local.get("last_update") or ""
                if cloud_lu and cloud_lu > local_lu:
                    payload = {k: v for k, v in crow.items()
                               if k not in ("id", "created", "updated", "inquiries")}
                    payload["cloud_id"] = cloud_id
                    payload["sync_status"] = "synced"
                    payload["local_updated_at"] = datetime.datetime.now(datetime.timezone.utc).isoformat()
                    # Filter only valid local columns
                    valid_cols = {"week", "due_date", "customer", "project_name", "sso_modeln",
                                  "ear", "application", "bu", "task_summary", "mchp_device",
                                  "project_schedule", "status", "category", "worklogs", "create_date",
                                  "last_update", "archive", "record_hash", "cloud_id", "sync_status",
                                  "local_updated_at"}
                    payload = {k: v for k, v in payload.items() if k in valid_cols}
                    sets = ", ".join([f"{k}=?" for k in payload])
                    c = db._conn()
                    # Guard: don't overwrite a soft-deleted record (race with concurrent delete)
                    c.execute(f"UPDATE worklog SET {sets} WHERE id=? AND sync_status != 'deleted'",
                              list(payload.values()) + [local["id"]])
                    c.commit(); c.close()
                    report["pulled"] += 1

    # ── PostgreSQL sync methods ───────────────────────────────────────────────

    _PG_COLS = frozenset({
        "week", "due_date", "customer", "project_name", "sso_modeln",
        "ear", "application", "bu", "task_summary", "mchp_device",
        "project_schedule", "todo_content", "todo_due_date",
        "status", "category", "worklogs", "create_date", "last_update",
        "archived", "record_hash",
    })

    _LOCAL_COLS = frozenset({
        "week", "due_date", "customer", "project_name", "sso_modeln",
        "ear", "application", "bu", "task_summary", "mchp_device",
        "project_schedule", "todo_content", "todo_due_date",
        "status", "category", "worklogs", "create_date", "last_update",
        "archive", "record_hash", "cloud_id", "sync_status", "local_updated_at",
    })

    def _push_postgres(self, report: dict):
        """Push pending local rows to PostgreSQL."""
        db = _get_local_db()
        pending = db.pending_rows()
        if not pending:
            return

        conn = _postgres_client()
        if not conn:
            report["errors"].append("PostgreSQL connection failed")
            return

        try:
            cur = conn.cursor()

            for row in pending:
                local_id = row["id"]
                cloud_id = row.get("cloud_id")

                if row["sync_status"] == "deleted":
                    pg_deleted = True
                    if cloud_id:
                        try:
                            pg_id = int(cloud_id)
                            cur.execute("DELETE FROM worklog WHERE id=%s", (pg_id,))
                            conn.commit()
                        except (ValueError, TypeError):
                            pass  # Non-integer cloud_id (not a PG record) — delete locally only
                        except Exception as e:
                            conn.rollback()
                            report["errors"].append(f"pg delete {local_id}: {e}")
                            pg_deleted = False
                    if pg_deleted:
                        c = db._conn()
                        c.execute("DELETE FROM worklog WHERE id=?", (local_id,))
                        c.commit(); c.close()
                        report["pushed"] += 1
                    continue

                # Build payload: map local "archive" → pg "archived"
                payload = {}
                for k, v in row.items():
                    if k in ("id", "cloud_id", "sync_status", "local_updated_at"):
                        continue
                    if k == "archive":
                        # Ensure archive value is always 'Yes' or 'No'
                        payload["archived"] = "Yes" if v == "Yes" else "No"
                    elif k in self._PG_COLS:
                        payload[k] = v

                if cloud_id:
                    # Conflict check
                    try:
                        cur.execute("SELECT last_update FROM worklog WHERE id=%s", (int(cloud_id),))
                        pg_row = cur.fetchone()
                        if pg_row:
                            cloud_lu = pg_row["last_update"] or ""
                            local_lu = row.get("last_update") or ""
                            if cloud_lu and local_lu and cloud_lu > local_lu:
                                cur.execute("SELECT * FROM worklog WHERE id=%s", (int(cloud_id),))
                                cloud_data = dict(cur.fetchone())
                                # Normalize archived value
                                archived_val = cloud_data.pop("archived", None)
                                cloud_data["archive"] = "Yes" if archived_val == "Yes" else "No"
                                db.mark_conflict(local_id, cloud_data)
                                report["conflicts"] += 1
                                continue
                    except Exception:
                        pass
                    try:
                        sets = ", ".join([f"{k}=%s" for k in payload])
                        vals = list(payload.values()) + [int(cloud_id)]
                        cur.execute(f"UPDATE worklog SET {sets} WHERE id=%s", vals)
                        conn.commit()
                        db.mark_synced(local_id, cloud_id)
                        report["pushed"] += 1
                    except Exception as e:
                        conn.rollback()
                        report["errors"].append(f"pg update {local_id}: {e}")
                else:
                    try:
                        cols = ", ".join(payload.keys())
                        phs  = ", ".join(["%s"] * len(payload))
                        cur.execute(
                            f"INSERT INTO worklog ({cols}) VALUES ({phs}) RETURNING id",
                            list(payload.values())
                        )
                        conn.commit()
                        new_id = cur.fetchone()["id"]
                        db.mark_synced(local_id, str(new_id))
                        report["pushed"] += 1
                    except Exception as e:
                        conn.rollback()
                        report["errors"].append(f"pg insert {local_id}: {e}")
        finally:
            conn.close()

    def _pull_postgres(self, report: dict):
        """Pull all PostgreSQL rows into local cache."""
        db = _get_local_db()

        conn = _postgres_client()
        if not conn:
            report["errors"].append("PostgreSQL connection failed")
            return
        try:
            cur = conn.cursor()
            cur.execute("SELECT * FROM worklog")
            cloud_rows = [dict(r) for r in cur.fetchall()]
        except Exception as e:
            report["errors"].append(f"pg pull fetch: {e}")
            return
        finally:
            conn.close()

        now_iso = datetime.datetime.now(datetime.timezone.utc).isoformat()
        for crow in cloud_rows:
            cloud_id = str(crow["id"])
            # Map pg "archived" → local "archive"
            archived_val = crow.pop("archived", None)
            # Normalize: only 'Yes' or 'No', handle NULL/empty/any other value as 'No'
            crow["archive"] = "Yes" if archived_val == "Yes" else "No"
            for col in ("id", "created_at"):
                crow.pop(col, None)

            local = db.find_by_cloud_id(cloud_id)
            if local is None:
                payload = {k: v for k, v in crow.items() if k in self._LOCAL_COLS}
                payload["cloud_id"]        = cloud_id
                payload["sync_status"]     = "synced"
                payload["local_updated_at"] = now_iso
                try:
                    db.insert(payload)
                    report["pulled"] += 1
                except Exception as e:
                    report["errors"].append(f"pg pull insert {cloud_id}: {e}")
            elif local["sync_status"] == "synced":
                cloud_lu = crow.get("last_update") or ""
                local_lu = local.get("last_update") or ""
                if cloud_lu and cloud_lu > local_lu:
                    payload = {k: v for k, v in crow.items() if k in self._LOCAL_COLS}
                    payload["cloud_id"]        = cloud_id
                    payload["sync_status"]     = "synced"
                    payload["local_updated_at"] = now_iso
                    sets = ", ".join([f"{k}=?" for k in payload])
                    c = db._conn()
                    # Guard: don't overwrite a soft-deleted record (race with concurrent delete)
                    c.execute(f"UPDATE worklog SET {sets} WHERE id=? AND sync_status != 'deleted'",
                              list(payload.values()) + [local["id"]])
                    c.commit(); c.close()
                    report["pulled"] += 1


# ── Global sync engine — always starts; silently skips cycles when offline ────
_sync_engine = SyncEngine()

def _trigger_sync():
    """Trigger an immediate background sync on data change (non-blocking).
    Always spawns a thread; sync_now() handles the online check internally.
    Only spawns if a cloud backend is configured (avoids noise in local-only mode).
    """
    if _postgres_configured() or _pocketbase_configured():
        _threading.Thread(target=_sync_engine.sync_now, daemon=True).start()

def get_db():
    """Always return the local DB. Sync runs in background."""
    return _get_local_db()


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def make_hash(customer, project):
    key = f"{(customer or '').strip().lower()}|||{(project or '').strip().lower()}"
    return hashlib.md5(key.encode()).hexdigest()

def week_number(date_str):
    try:
        d = datetime.datetime.strptime(date_str, "%m-%d-%Y")
        return str(d.isocalendar()[1])
    except Exception:
        return ""

def today_str():
    return datetime.date.today().strftime("%m-%d-%Y")

def parse_worklog_entries(text):
    if not text: return []
    # Support both old `── YYYY-MM-DD HH:MM ──` and new `── YYYY-MM-DD ──`
    pattern = re.compile(r'── (\d{4}-\d{2}-\d{2})(?: \d{2}:\d{2})? ──')
    parts = pattern.split(text.strip())
    entries = []
    i = 1
    while i + 1 < len(parts):
        try:
            dt = datetime.datetime.strptime(parts[i].strip(), "%Y-%m-%d")
        except Exception:
            dt = datetime.datetime.min
        if parts[i+1].strip():
            entries.append((dt, parts[i+1].strip()))
        i += 2
    return entries

def merge_worklogs(a, b):
    all_e = parse_worklog_entries(a) + parse_worklog_entries(b)
    seen, unique = set(), []
    for dt, c in all_e:
        if c.strip() not in seen:
            seen.add(c.strip()); unique.append((dt, c))
    unique.sort(key=lambda x: x[0])
    return "\n\n".join(f"── {dt.strftime('%Y-%m-%d')} ──\n{c}" for dt, c in unique)


# ─────────────────────────────────────────────────────────────────────────────
# API ROUTES
# ─────────────────────────────────────────────────────────────────────────────
@app.route("/api/records", methods=["GET"])
def api_list():
    filters = {k: request.args.get(k,"") for k in
               ("customer","project_name","status","category","archive")}
    return jsonify(get_db().fetch_all(filters))

@app.route("/api/records", methods=["POST"])
def api_create():
    try:
        data = request.json
        if not data:
            return jsonify({"error": "No data received"}), 400
        data["week"]        = week_number(data.get("last_update","") or today_str())
        data["record_hash"] = make_hash(data.get("customer",""), data.get("project_name",""))
        data.setdefault("create_date", today_str())
        data.setdefault("last_update", today_str())
        result = get_db().insert(data)
        return jsonify(result), 201
    except Exception as e:
        app.logger.error(f"Insert error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/records/<int:rid>", methods=["PUT"])
def api_update(rid):
    try:
        data = request.json
        data["week"]        = week_number(data.get("last_update","") or today_str())
        data["record_hash"] = make_hash(data.get("customer",""), data.get("project_name",""))
        result = get_db().update(rid, data)
        return jsonify(result)
    except Exception as e:
        app.logger.error(f"Update error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/debug/db")
def api_debug():
    """Quick health check — shows row count and DB path."""
    try:
        db = get_db()
        rows = db.fetch_all()
        return jsonify({"ok": True, "count": len(rows), "type": type(db).__name__,
                        "path": getattr(db, "path", "cloud")})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/api/records/<int:rid>", methods=["DELETE"])
def api_delete(rid):
    get_db().delete(rid)
    return jsonify({"ok": True})

@app.route("/api/import", methods=["POST"])
def api_import():
    """Import all records from an uploaded SQLite database file."""
    if 'db_file' not in request.files:
        return jsonify({"error": "未選擇檔案"}), 400

    file = request.files['db_file']
    if file.filename == '':
        return jsonify({"error": "檔案名稱為空"}), 400

    if not file.filename.endswith(('.db', '.sqlite', '.sqlite3')):
        return jsonify({"error": "檔案格式錯誤，僅支援 .db、.sqlite、.sqlite3"}), 400

    temp_fd, temp_path = tempfile.mkstemp(suffix='.db')

    try:
        # Save uploaded file to temporary location
        os.close(temp_fd)
        file.save(temp_path)

        # Connect to the uploaded database
        conn = sqlite3.connect(temp_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()

        # Check if worklog table exists
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='worklog'")
        if not cursor.fetchone():
            conn.close()
            return jsonify({"error": "資料庫中找不到 'worklog' 資料表"}), 400

        # Fetch all records from the uploaded database
        cursor.execute("SELECT * FROM worklog")
        rows = cursor.fetchall()
        conn.close()

        # Import records into local database
        db = get_db()
        added = merged = skipped = 0
        results = []

        for row in rows:
            # Convert sqlite3.Row to dict
            record = dict(row)

            # Normalize field names and values
            normalized = {}
            for field in ["week", "due_date", "customer", "project_name", "sso_modeln",
                          "ear", "application", "bu", "task_summary", "mchp_device",
                          "project_schedule", "todo_content", "todo_due_date", "status",
                          "category", "worklogs", "create_date", "last_update"]:
                normalized[field] = record.get(field, "") or ""

            # Handle different archive field names
            if "archive" in record:
                normalized["archive"] = record["archive"] or "No"
            elif "archieved" in record:
                normalized["archive"] = record["archieved"] or "No"
            else:
                normalized["archive"] = "No"

            # Handle inquiries → mchp_device mapping for old databases
            if "inquiries" in record and not normalized["mchp_device"]:
                normalized["mchp_device"] = record["inquiries"] or ""

            # Check if record exists (by customer + project_name)
            h = make_hash(normalized.get("customer", ""), normalized.get("project_name", ""))
            existing = db.find_by_hash(h)

            if existing is None:
                # New record - add it
                normalized["record_hash"] = h
                # Remove any extra fields not in schema
                for key in list(normalized.keys()):
                    if key not in ["week", "due_date", "customer", "project_name", "sso_modeln",
                                   "ear", "application", "bu", "task_summary", "mchp_device",
                                   "project_schedule", "todo_content", "todo_due_date", "status",
                                   "category", "worklogs", "create_date", "last_update", "archive",
                                   "record_hash"]:
                        normalized.pop(key, None)
                db.insert(normalized)
                added += 1
                results.append({
                    "status": "added",
                    "customer": normalized.get("customer", ""),
                    "project": normalized.get("project_name", "")
                })
            else:
                # Existing record - merge worklogs
                mwl = merge_worklogs(existing.get("worklogs", ""), normalized.get("worklogs", ""))
                if mwl != (existing.get("worklogs", "") or ""):
                    upd = dict(existing)
                    eid = upd.pop("id")
                    upd["worklogs"] = mwl
                    upd["last_update"] = today_str()
                    db.update(eid, upd)
                    merged += 1
                    results.append({
                        "status": "merged",
                        "customer": normalized.get("customer", ""),
                        "project": normalized.get("project_name", "")
                    })
                else:
                    skipped += 1
                    results.append({
                        "status": "skipped",
                        "customer": normalized.get("customer", ""),
                        "project": normalized.get("project_name", "")
                    })

        return jsonify({
            "added": added,
            "merged": merged,
            "skipped": skipped,
            "results": results
        })

    except sqlite3.Error as e:
        return jsonify({"error": f"讀取資料庫錯誤: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": f"匯入過程發生錯誤: {str(e)}"}), 500
    finally:
        # Clean up temporary file
        try:
            if os.path.exists(temp_path):
                os.unlink(temp_path)
        except:
            pass

@app.route("/api/sync/init", methods=["GET"])
def api_sync_init():
    """Called by frontend on page load — triggers initial pull if local DB is empty."""
    has_data = _has_local_data()
    online   = _is_online()
    return jsonify({
        "has_local_data": has_data,
        "online":         online,
        "db_path":        SQLITE_PATH,
    })

@app.route("/api/sync/status", methods=["GET"])
def api_sync_status():
    return jsonify(_sync_engine.status())

@app.route("/api/sync/test", methods=["GET"])
def api_sync_test():
    """Verify connectivity and return detailed diagnostics."""
    if _postgres_configured():
        url = os.environ.get("POSTGRES_URL", POSTGRES_URL)
        pg_error = None
        online = False
        try:
            conn = psycopg2.connect(url, cursor_factory=RealDictCursor,
                                    connect_timeout=5)
            conn.close()
            online = True
        except Exception as e:
            pg_error = str(e).split("\n")[0]  # first line only
        short_url = url[:60] + "…" if len(url) > 60 else url
        return jsonify({
            "configured":        True,
            "online":            online,
            "url":               short_url,
            "has_pocketbase_lib": HAS_POCKETBASE,
            "backend":           "postgres",
            "error":             pg_error,
        })
    configured = _pocketbase_configured()
    online     = _is_online() if configured else False
    url        = os.environ.get("POCKETBASE_URL", POCKETBASE_URL)
    return jsonify({
        "configured":        configured,
        "online":            online,
        "url":               url[:40] + "…" if len(url) > 40 else url,
        "has_pocketbase_lib": HAS_POCKETBASE,
        "backend":           "pocketbase",
    })

@app.route("/api/sync/now", methods=["POST"])
def api_sync_now():
    """Trigger an immediate sync attempt."""
    report = _sync_engine.sync_now()
    return jsonify(report)

@app.route("/api/sync/conflicts", methods=["GET"])
def api_sync_conflicts():
    """Return all unresolved conflicts."""
    conflicts = _get_local_db().get_conflicts()
    return jsonify(conflicts)

@app.route("/api/sync/conflicts/<int:conflict_id>", methods=["POST"])
def api_sync_resolve(conflict_id):
    """
    Resolve a conflict.
    Body: { "keep": "local" | "cloud" | "backup" }
    """
    keep = (request.json or {}).get("keep", "local")
    if keep not in ("local", "cloud", "backup"):
        return jsonify({"error": "keep must be 'local', 'cloud', or 'backup'"}), 400
    _get_local_db().resolve_conflict(conflict_id, keep)
    # If keeping local or backup, trigger immediate sync push
    if keep in ("local", "backup"):
        _threading.Thread(target=_sync_engine.sync_now, daemon=True).start()
    return jsonify({"ok": True, "kept": keep})

@app.route("/api/config", methods=["GET"])
def api_config():
    sync = _sync_engine.status()
    return jsonify({
        "status_options":   STATUS_OPTIONS,
        "category_options": CATEGORY_OPTIONS,
        "today":            today_str(),
        "db_type":          "PostgreSQL ☁" if _postgres_configured() else ("PocketBase + Local Cache" if _pocketbase_configured() else "SQLite (local only)"),
        "has_docx":         HAS_DOCX,
        "version":          APP_VERSION,
        "sync":             sync,
    })

# Always start the sync engine — it checks connectivity internally each cycle
_sync_engine.start()


@app.route("/api/export/weeks", methods=["POST"])
def api_export_weeks():
    """Return available weeks from selected records' worklogs."""
    body = request.json or {}
    ids  = body.get("ids", [])
    all_rows = get_db().fetch_all()

    if ids:
        id_set = set(int(i) for i in ids)
        rows = [r for r in all_rows if int(r.get("id", 0)) in id_set]
    else:
        rows = all_rows

    # Collect all unique (year, week) from worklog entries
    weeks_set = set()
    for row in rows:
        for dt, _ in parse_worklog_entries(row.get("worklogs") or ""):
            if dt.year > 1970:
                weeks_set.add((dt.isocalendar()[0], dt.isocalendar()[1]))

    # Sort descending (newest first)
    weeks_list = sorted(weeks_set, reverse=True)

    # Build response with date ranges
    result = []
    for year, week in weeks_list:
        mon, fri = _week_date_range(year, week)
        result.append({
            "year": year,
            "week": week,
            "label": f"Week {week}, {year}",
            "range": f"{mon.strftime('%m/%d')} – {fri.strftime('%m/%d')}",
            "value": f"{year}-W{week:02d}"
        })

    return jsonify(result)


@app.route("/api/export/word", methods=["POST"])
def api_export_word():
    if not HAS_DOCX:
        return jsonify({"error": "python-docx not installed — run: pip install python-docx"}), 500
    try:
        body     = request.json or {}
        ids      = body.get("ids", [])       # list of record IDs to export
        week_str = body.get("week", "")      # optional: "YYYY-WNN" format
        all_rows = get_db().fetch_all()

        # Filter to selected IDs only
        if ids:
            id_set = set(int(i) for i in ids)
            rows   = [r for r in all_rows if int(r.get("id", 0)) in id_set]
        else:
            rows = all_rows

        if not rows:
            return jsonify({"error": "No records match the selection"}), 400

        # Parse week parameter if provided
        target_year, target_week = None, None
        if week_str and week_str != "all":
            try:
                # Format: "YYYY-WNN"
                parts = week_str.split("-W")
                target_year = int(parts[0])
                target_week = int(parts[1])
            except (ValueError, IndexError):
                pass

        buf = io.BytesIO()
        _build_word_doc(rows, buf, target_year=target_year, target_week=target_week)
        buf.seek(0)
        if buf.getbuffer().nbytes < 100:
            return jsonify({"error": "Document generation produced empty output"}), 500

        fname = f"WorkLog_{datetime.date.today().strftime('%Y%m%d')}.docx"
        return send_file(buf, as_attachment=True, download_name=fname,
                         mimetype="application/vnd.openxmlformats-officedocument"
                                   ".wordprocessingml.document")
    except Exception as e:
        app.logger.error(f"Word export error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


def _export_excel_wlr(rows):
    """Export WLR format (v0.9.5 comprehensive format)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "WorkLog Report"

    rows = [r for r in rows if (r.get("archive") or "No") != "Yes"]

    if not rows:
        return jsonify({"error": "No records to export (all selected records are archived)"}), 400

    # Column mapping: DB field -> Excel header
    columns = [
        ("customer", "OEM/ODM/OBM/Disti"),
        ("sso_modeln", "SSO#"),
        ("project_name", "Project Name"),
        ("application", "Application"),
        ("bu", "BU"),
        ("mchp_device", "Microchip Devices"),
        ("task_summary", "Status/Update"),
        ("ear", "EAR(K$)"),
        ("project_schedule", "Timeline"),
    ]

    # Header style
    header_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    # Write headers
    for col_idx, (_, header) in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # Write data
    for row_idx, row in enumerate(rows, 2):
        for col_idx, (field, _) in enumerate(columns, 1):
            value = row.get(field, "") or ""
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(vertical='top', wrap_text=True)

    # Auto-adjust column widths
    for col_idx, (_, header) in enumerate(columns, 1):
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = 15

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    fname = f"WorkLog_WLR_{datetime.date.today().strftime('%Y%m%d')}.xlsx"
    return send_file(buf, as_attachment=True, download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def _export_excel_todo(rows, date_filter):
    """Export To-Do format with checkboxes and date filtering."""
    from datetime import datetime as dt, timedelta

    # Filter: must have non-empty todo_due_date
    filtered_rows = []
    for row in rows:
        due_date = (row.get("todo_due_date") or "").strip()
        if not due_date:
            continue

        # Apply Week+1 filter if requested
        if date_filter == "week_plus_1":
            try:
                # Parse MM-DD-YYYY format
                parts = due_date.split("-")
                if len(parts) == 3:
                    month, day, year = int(parts[0]), int(parts[1]), int(parts[2])
                    due_date_obj = dt(year, month, day)
                    today = dt.now()
                    week_from_now = today + timedelta(days=7)

                    if today <= due_date_obj <= week_from_now:
                        filtered_rows.append(row)
            except (ValueError, IndexError):
                continue
        else:
            filtered_rows.append(row)

    if not filtered_rows:
        return jsonify({"error": "No To-Do records with due dates found"}), 400

    # Sort by due date ascending
    def parse_date_for_sort(row):
        due_date = row.get("todo_due_date") or ""
        try:
            parts = due_date.split("-")
            if len(parts) == 3:
                return dt(int(parts[2]), int(parts[0]), int(parts[1]))
        except:
            pass
        return dt(1970, 1, 1)

    filtered_rows.sort(key=parse_date_for_sort)

    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "To-Do List"

    # Columns: Completed (☐), Customer, Project Name, Due Date, To-Do
    columns = [
        ("completed", "Completed"),
        ("customer", "Customer"),
        ("project_name", "Project Name"),
        ("todo_due_date", "Due Date"),
        ("todo_content", "To-Do"),
    ]

    # Header styling
    header_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    # Write headers
    for col_idx, (_, header) in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # Write data
    for row_idx, row in enumerate(filtered_rows, 2):
        for col_idx, (field, _) in enumerate(columns, 1):
            if field == "completed":
                value = "☐"  # Unicode unchecked box
            else:
                value = row.get(field, "") or ""

            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border

            if field == "completed":
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.font = Font(size=14)  # Larger checkbox
            elif field == "todo_content":
                cell.alignment = Alignment(vertical='top', wrap_text=True)
            else:
                cell.alignment = Alignment(vertical='center')

    # Column widths
    ws.column_dimensions['A'].width = 10   # Completed
    ws.column_dimensions['B'].width = 20   # Customer
    ws.column_dimensions['C'].width = 25   # Project Name
    ws.column_dimensions['D'].width = 12   # Due Date
    ws.column_dimensions['E'].width = 40   # To-Do

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    fname = f"WorkLog_ToDo_{datetime.date.today().strftime('%Y%m%d')}.xlsx"
    return send_file(buf, as_attachment=True, download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.route("/api/export/excel", methods=["POST"])
def api_export_excel():
    """Export selected records to Excel with format options."""
    if not HAS_EXCEL:
        return jsonify({"error": "openpyxl not installed — run: pip install openpyxl"}), 500
    try:
        body = request.json or {}
        ids = body.get("ids", [])
        export_format = body.get("format", "WLR")
        date_filter = body.get("date_filter", "all")

        all_rows = get_db().fetch_all()

        if ids:
            id_set = set(int(i) for i in ids)
            rows = [r for r in all_rows if int(r.get("id", 0)) in id_set]
        else:
            rows = all_rows

        if not rows:
            return jsonify({"error": "No records to export"}), 400

        if export_format == "TODO":
            return _export_excel_todo(rows, date_filter)
        else:
            return _export_excel_wlr(rows)

    except Exception as e:
        app.logger.error(f"Excel export error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route("/api/export/pdf", methods=["POST"])
def api_export_pdf():
    """Export selected records to PDF."""
    if not HAS_PDF:
        return jsonify({"error": "fpdf2 not installed — run: pip install fpdf2"}), 500
    try:
        body = request.json or {}
        ids = body.get("ids", [])
        all_rows = get_db().fetch_all()

        if ids:
            id_set = set(int(i) for i in ids)
            rows = [r for r in all_rows if int(r.get("id", 0)) in id_set]
        else:
            rows = all_rows

        if not rows:
            return jsonify({"error": "No records to export"}), 400

        pdf = FPDF(orientation='L', unit='mm', format='A4')
        pdf.set_auto_page_break(auto=True, margin=10)
        pdf.add_page()

        # Calculate available width (A4 landscape = 297mm, minus margins)
        page_width = 297 - 20  # 277mm usable width

        # Title
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "Work Log Report", ln=True, align='C')
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 5, f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  |  {len(rows)} record(s)", ln=True, align='C')
        pdf.ln(4)

        # Column setup for A4 Landscape - optimized to fit page width
        headers = ["Customer", "SSO#", "Project", "Application", "BU", "MCHP Device", "Task Summary", "EAR(K$)", "Timeline"]
        col_widths = [28, 15, 26, 30, 12, 28, 60, 15, 25]  # Total = 239mm

        # Auto-scale to fit page width
        total_width = sum(col_widths)
        if total_width > page_width:
            scale_factor = page_width / total_width
            col_widths = [w * scale_factor for w in col_widths]

        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(30, 58, 95)
        pdf.set_text_color(255, 255, 255)
        for i, header in enumerate(headers):
            pdf.cell(col_widths[i], 7, header, border=1, align='C', fill=True)
        pdf.ln()

        # Data rows with multi-line support for Task Summary column
        pdf.set_font("Helvetica", "", 7)
        pdf.set_text_color(0, 0, 0)

        for row in rows:
            # Prepare data
            customer = (row.get("customer") or "")
            sso = (row.get("sso_modeln") or "")
            project = (row.get("project_name") or "")
            application = (row.get("application") or "")
            bu = (row.get("bu") or "")
            device = (row.get("mchp_device") or "")
            task_summary = (row.get("task_summary") or "")  # Allow full text with wrapping
            ear = (row.get("ear") or "")
            timeline = (row.get("project_schedule") or "")

            # Calculate row height based on Task Summary text (allow wrapping)
            summary_width_mm = col_widths[6]
            chars_per_line = int(summary_width_mm / 1.3)  # Approx chars per line at 7pt

            # Split Task Summary into lines (word wrap)
            summary_lines = []
            if task_summary:
                words = task_summary.split()
                current_line = ""
                for word in words:
                    test_line = current_line + (" " if current_line else "") + word
                    if len(test_line) <= chars_per_line:
                        current_line = test_line
                    else:
                        if current_line:
                            summary_lines.append(current_line)
                        current_line = word
                if current_line:
                    summary_lines.append(current_line)

            if not summary_lines:
                summary_lines = [""]

            # Calculate row height
            line_height = 4
            row_height = max(6, len(summary_lines) * line_height)

            # Store starting position
            start_x = pdf.get_x()
            start_y = pdf.get_y()

            # Draw cells - first 6 columns (before Task Summary)
            data_cols = [customer, sso, project, application, bu, device]
            x_offset = 0
            for i, val in enumerate(data_cols):
                # Truncate to fit
                max_chars = int(col_widths[i] / 1.3)
                truncated = val[:max_chars] if len(val) > max_chars else val
                pdf.set_xy(start_x + x_offset, start_y)
                pdf.cell(col_widths[i], row_height, truncated, border=1, align='L')
                x_offset += col_widths[i]

            # Draw Task Summary column with multi-line text
            summary_x = start_x + x_offset
            pdf.rect(summary_x, start_y, col_widths[6], row_height)
            for line_idx, line in enumerate(summary_lines):
                pdf.set_xy(summary_x + 1, start_y + 1 + line_idx * line_height)
                pdf.cell(col_widths[6] - 2, line_height, line, border=0, align='L')
            x_offset += col_widths[6]

            # Draw last 2 columns (EAR, Timeline)
            pdf.set_xy(start_x + x_offset, start_y)
            ear_truncated = ear[:int(col_widths[7] / 1.3)]
            pdf.cell(col_widths[7], row_height, ear_truncated, border=1, align='L')
            x_offset += col_widths[7]
            pdf.set_xy(start_x + x_offset, start_y)
            timeline_truncated = timeline[:int(col_widths[8] / 1.3)]
            pdf.cell(col_widths[8], row_height, timeline_truncated, border=1, align='L')

            # Move to next row
            pdf.set_xy(start_x, start_y + row_height)

        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)

        fname = f"WorkLog_{datetime.date.today().strftime('%Y%m%d')}.pdf"
        return send_file(buf, as_attachment=True, download_name=fname, mimetype="application/pdf")
    except Exception as e:
        app.logger.error(f"PDF export error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


_db_size_cache: dict = {"size": None, "bytes": 0, "ts": 0.0}
_DB_SIZE_TTL = 60  # seconds

@app.route("/api/db/size", methods=["GET"])
def api_db_size():
    """Return the actual WorkLog.db file size, cached for 60 s."""
    import time
    now = time.monotonic()
    if _db_size_cache["size"] is not None and now - _db_size_cache["ts"] < _DB_SIZE_TTL:
        return jsonify({"size": _db_size_cache["size"], "bytes": _db_size_cache["bytes"]})
    try:
        size_bytes = os.path.getsize(SQLITE_PATH)
        if size_bytes < 1024:
            size_str = f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            size_str = f"{size_bytes / 1024:.1f} KB"
        else:
            size_str = f"{size_bytes / (1024 * 1024):.2f} MB"
        _db_size_cache.update({"size": size_str, "bytes": size_bytes, "ts": now})
        return jsonify({"size": size_str, "bytes": size_bytes})
    except Exception as e:
        return jsonify({"size": "N/A", "bytes": 0, "error": str(e)})


@app.route("/api/sync/check-schema", methods=["GET"])
def api_sync_check_schema():
    """Check if PocketBase database has all required columns."""
    if not _pocketbase_configured():
        return jsonify({"configured": False, "message": "PocketBase not configured"})

    try:
        pb = _cloud_client()
        # Try to fetch one row to check available columns
        records = pb.collection(TABLE).get_list(1, 1)

        required_cols = {"sso_modeln", "ear", "bu", "application", "mchp_device", "project_schedule",
                         "todo_content", "todo_due_date"}
        if records.items:
            existing_cols = set(vars(records.items[0]).keys())
            missing = required_cols - existing_cols
            if missing:
                return jsonify({
                    "ok": False,
                    "missing_columns": list(missing),
                    "message": f"PocketBase schema needs update. Missing: {', '.join(missing)}"
                })
        return jsonify({"ok": True, "message": "Schema is up to date"})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/env", methods=["GET"])
def api_env_get():
    """Return current sync configuration (token value redacted)."""
    token = os.environ.get("POCKETBASE_TOKEN", POCKETBASE_TOKEN)
    return jsonify({
        "POCKETBASE_URL":   os.environ.get("POCKETBASE_URL", POCKETBASE_URL),
        "POCKETBASE_TOKEN": "••••" if token else "",
        "POSTGRES_URL":     os.environ.get("POSTGRES_URL", POSTGRES_URL),
    })


@app.route("/api/env", methods=["POST"])
def api_env_set():
    """Save sync configuration to .env and reload env vars in-process."""
    global POCKETBASE_URL, POCKETBASE_TOKEN, POSTGRES_URL
    data = request.json or {}

    env_path = Path(__file__).parent / ".env"
    existing: dict = {}
    if env_path.exists():
        for line in env_path.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, _, v = line.partition("=")
                existing[k.strip()] = v.strip()

    for key in ("POCKETBASE_URL", "POCKETBASE_TOKEN", "POSTGRES_URL"):
        if key in data and data[key] != "••••":
            existing[key] = data[key].strip()

    lines = ["# Work Log Journal — Environment Variables", ""]
    lines += [f"{k}={v}" for k, v in existing.items()]
    env_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    if "POCKETBASE_URL" in data:
        POCKETBASE_URL = existing.get("POCKETBASE_URL", "")
        os.environ["POCKETBASE_URL"] = POCKETBASE_URL
    if "POCKETBASE_TOKEN" in data and data["POCKETBASE_TOKEN"] != "••••":
        POCKETBASE_TOKEN = existing.get("POCKETBASE_TOKEN", "")
        os.environ["POCKETBASE_TOKEN"] = POCKETBASE_TOKEN
    if "POSTGRES_URL" in data:
        POSTGRES_URL = existing.get("POSTGRES_URL", "")
        os.environ["POSTGRES_URL"] = POSTGRES_URL

    return jsonify({"ok": True})


@app.route("/api/image/resize", methods=["POST"])
def api_image_resize():
    """Accept base64 PNG, return resized thumbnail base64."""
    if not HAS_PIL:
        return jsonify({"error": "Pillow not installed"}), 500
    b64 = request.json.get("data", "")
    try:
        raw  = base64.b64decode(b64)
        img  = Image.open(io.BytesIO(raw))
        img.thumbnail((800, 600), Image.LANCZOS)
        out  = io.BytesIO()
        img.save(out, format="PNG")
        return jsonify({"data": base64.b64encode(out.getvalue()).decode()})
    except Exception as e:
        return jsonify({"error": str(e)}), 400


# ─────────────────────────────────────────────────────────────────────────────
# WORD BUILDER
# ─────────────────────────────────────────────────────────────────────────────
def _hex_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def _shade(cell, hx):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), hx.lstrip("#")); tcPr.append(shd)

def _cell(cell, text, bold=False, sz=10, bg=None, align="left"):
    cell.text=""
    p=cell.paragraphs[0]
    p.alignment={"center":WD_ALIGN_PARAGRAPH.CENTER}.get(align,WD_ALIGN_PARAGRAPH.LEFT)
    run=p.add_run(str(text or "")); run.font.size=Pt(sz); run.font.bold=bold
    if bg: _shade(cell,bg)

def _hdr_row(row, headers, bg="1E3A5F", sz=10):
    for cell,h in zip(row.cells, headers):
        _shade(cell,bg); p=cell.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(h); run.font.size=Pt(sz); run.font.bold=True
        run.font.color.rgb=_hex_rgb("FFFFFF")

def _set_widths(table, widths):
    # Set column widths (applies to all rows including future ones)
    for col, w in zip(table.columns, widths):
        col.width = w
    # Also set cell widths for existing rows
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

def _week_date_range(year, week):
    """Return (monday, friday) as date objects for ISO week."""
    jan4      = datetime.date(year, 1, 4)
    week1_mon = jan4 - datetime.timedelta(days=jan4.weekday())
    monday    = week1_mon + datetime.timedelta(weeks=week - 1)
    friday    = monday + datetime.timedelta(days=4)
    return monday, friday


def _wl_content_to_word(cell, content, bg_hex, doc):
    """
    Render a Worklog entry body into a Word table cell.
    Handles: Markdown bold/italic/code, headings, bullets, indent, images.
    """
    IMG_PAT   = re.compile(r'\[IMG_B64:([A-Za-z0-9+/=]+)\]')
    STAMP_PAT = re.compile(r'^── \d{4}-\d{2}-\d{2}(?: \d{2}:\d{2})? ──$')
    LIST_PAT  = re.compile(r'^([ \t]*)([-*+]|\d+\.)\s+(.*)')
    HDR_PAT   = re.compile(r'^(#{1,3})\s+(.*)')

    def add_inline(para, text, base_sz=10):
        # Pattern: bold(**), italic(*), underline(__), strikethrough(~~), code(`)
        for seg in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|~~[^~]+~~|`[^`]+`)', text):
            if not seg: continue
            run = para.add_run()
            run.font.size = Pt(base_sz)
            if seg.startswith('**') and seg.endswith('**'):
                run.text = seg[2:-2]; run.font.bold = True
            elif seg.startswith('*') and seg.endswith('*'):
                run.text = seg[1:-1]; run.font.italic = True
            elif seg.startswith('__') and seg.endswith('__'):
                run.text = seg[2:-2]; run.font.underline = True
            elif seg.startswith('~~') and seg.endswith('~~'):
                run.text = seg[2:-2]; run.font.strike = True
            elif seg.startswith('`') and seg.endswith('`'):
                run.text = seg[1:-1]; run.font.name = 'Courier New'; run.font.size = Pt(8)
            else:
                run.text = seg

    cell.text = ''; _shade(cell, bg_hex)
    first_para = True

    def get_para():
        nonlocal first_para
        p = cell.paragraphs[0] if first_para else cell.add_paragraph()
        first_para = False
        _shade(cell, bg_hex)
        return p

    for line in content.split('\n'):
        # Pure image line
        img_m = IMG_PAT.fullmatch(line.strip())
        if img_m:
            try:
                p = get_para()
                p.add_run().add_picture(io.BytesIO(base64.b64decode(img_m.group(1))), width=Cm(8))
            except Exception:
                get_para().add_run('[image]').font.size = Pt(8)
            continue

        raw_indent = re.match(r'^([ \t]*)', line).group(1)
        indent_sp  = raw_indent.replace('\t', '  ')
        indent_lvl = len(indent_sp) // 2
        indent_cm  = Cm(indent_lvl * 0.63)
        stripped   = line.lstrip()

        if STAMP_PAT.match(stripped):
            p = get_para()
            norm = re.sub(r'── (\d{4}-\d{2}-\d{2})(?:(?: \d{2}:\d{2}))? ──', r'── \1 ──', stripped)
            run = p.add_run(norm); run.font.bold = True
            run.font.size = Pt(8); run.font.color.rgb = _hex_rgb("2E75B6")
            continue

        hm = HDR_PAT.match(stripped)
        if hm:
            p = get_para(); p.paragraph_format.left_indent = indent_cm
            run = p.add_run(hm.group(2)); run.font.bold = True
            run.font.size = Pt(10 - hm.group(1).count('#')); run.font.color.rgb = _hex_rgb("1E3A5F")
            continue

        lm = LIST_PAT.match(line)
        if lm:
            bullet = lm.group(2); text = lm.group(3)
            p = get_para()
            p.paragraph_format.left_indent = Cm((indent_lvl + 1) * 0.63)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            prefix = '•' if bullet in ('-', '*', '+') else bullet
            p.add_run(f'{prefix}  ').font.size = Pt(10)
            add_inline(p, text)
            continue

        segs = IMG_PAT.split(line)
        if len(segs) > 1:
            for si, seg in enumerate(segs):
                if si % 2 == 0:
                    if seg.strip():
                        p = get_para(); p.paragraph_format.left_indent = indent_cm
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        add_inline(p, seg.strip())
                else:
                    try:
                        p = get_para(); p.paragraph_format.left_indent = indent_cm
                        p.add_run().add_picture(io.BytesIO(base64.b64decode(seg)), width=Cm(8))
                    except Exception:
                        get_para().add_run('[image]').font.size = Pt(8)
            continue

        if stripped:
            p = get_para(); p.paragraph_format.left_indent = indent_cm
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, stripped)


def _build_word_doc(all_rows, out_buf, target_year=None, target_week=None):
    """Build Word report: Task Summary + selected week Worklogs with Markdown/image rendering."""
    # If no target week specified, use all records
    if target_year and target_week:
        wmon, wfri = _week_date_range(target_year, target_week)
        week_str = (f"Week {target_week}  ·  "
                    f"{wmon.strftime('%Y-%m-%d')} – {wfri.strftime('%Y-%m-%d')}")
    else:
        target_year = target_week = None
        week_str = "All records"

    doc = DocxDocument()
    try:
        cp = doc.core_properties
        cp.author = "Work Log Journal"; cp.title = "Work Log Report"
        cp.subject = "Work Log"; cp.keywords = "worklog"
    except Exception:
        pass
    settings_part = doc.settings.element
    for child in settings_part.findall(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}documentProtection'):
        settings_part.remove(child)

    sec = doc.sections[0]
    sec.page_width  = Cm(29.7); sec.page_height = Cm(21.0)
    sec.left_margin = sec.right_margin = Cm(1.8)
    sec.top_margin  = sec.bottom_margin = Cm(1.5)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Work Log Report")
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = _hex_rgb("1E3A5F")

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{week_str}   |   "
                  f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}   |   "
                  f"{len(all_rows)} record(s)")
    r.font.size = Pt(9); r.font.color.rgb = _hex_rgb("6B7280")
    doc.add_paragraph()

    report_rows = []
    for row in all_rows:
        wl = row.get("worklogs") or ""
        if target_year and target_week:
            # Filter entries for specific week
            entries = [(dt, c) for dt, c in parse_worklog_entries(wl)
                       if dt.isocalendar()[:2] == (target_year, target_week)]
            # Only include records that have entries in the target week
            if not entries:
                continue
        else:
            entries = parse_worklog_entries(wl)
            if not entries and wl.strip():
                entries = [(datetime.datetime(1970, 1, 1), wl.strip())]
            # For "all records", include if has entries or task_summary
            if not entries and not (row.get("task_summary") or "").strip():
                continue

        report_rows.append({
            "customer":     row.get("customer") or "",
            "project_name": row.get("project_name") or "",
            "task_summary": row.get("task_summary") or "",
            "entries":      sorted(entries, key=lambda x: x[0]),
        })
    report_rows.sort(key=lambda r: (r["customer"].lower(), r["project_name"].lower()))

    for idx, rec in enumerate(report_rows):
        hdr = doc.add_table(2, 2); hdr.style = "Table Grid"
        _set_widths(hdr, [Cm(3), Cm(22.9)])
        # Customer row
        _shade(hdr.rows[0].cells[0], "1E3A5F"); _shade(hdr.rows[0].cells[1], "1E3A5F")
        lbl0 = hdr.rows[0].cells[0].paragraphs[0].add_run("Customer")
        lbl0.font.size = Pt(10); lbl0.font.bold = True; lbl0.font.color.rgb = _hex_rgb("FFFFFF")
        hdr.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        val0 = hdr.rows[0].cells[1].paragraphs[0].add_run(rec['customer'])
        val0.font.size = Pt(10); val0.font.bold = True; val0.font.color.rgb = _hex_rgb("FFFFFF")
        # Project row
        _shade(hdr.rows[1].cells[0], "1E3A5F"); _shade(hdr.rows[1].cells[1], "1E3A5F")
        lbl1 = hdr.rows[1].cells[0].paragraphs[0].add_run("Project")
        lbl1.font.size = Pt(10); lbl1.font.bold = True; lbl1.font.color.rgb = _hex_rgb("FFFFFF")
        hdr.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        val1 = hdr.rows[1].cells[1].paragraphs[0].add_run(rec['project_name'])
        val1.font.size = Pt(10); val1.font.bold = True; val1.font.color.rgb = _hex_rgb("FFFFFF")

        ts = (rec["task_summary"] or "").strip()
        if ts:
            ts_tbl = doc.add_table(1, 2); ts_tbl.style = "Table Grid"
            _set_widths(ts_tbl, [Cm(3), Cm(22.9)])
            _shade(ts_tbl.rows[0].cells[0], "344B6E")
            lbl = ts_tbl.rows[0].cells[0].paragraphs[0].add_run("Task Summary")
            lbl.font.size = Pt(10); lbl.font.bold = True; lbl.font.color.rgb = _hex_rgb("FFFFFF")
            ts_tbl.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _wl_content_to_word(ts_tbl.rows[0].cells[1], ts, "F0F4FB", doc)

        if rec["entries"]:
            et = doc.add_table(1, 2); et.style = "Table Grid"
            _set_widths(et, [Cm(3), Cm(22.9)])
            _hdr_row(et.rows[0], ["Date", "Worklog Entry"], bg="344B6E")
            for i, (dt, content) in enumerate(rec["entries"]):
                er = et.add_row()
                bg = "EEF3FA" if i % 2 == 0 else "FFFFFF"
                date_str = dt.strftime("%Y-%m-%d") if dt.year != 1970 else "—"
                _cell(er.cells[0], date_str, sz=10, bg=bg, align="center")
                _wl_content_to_word(er.cells[1], content, bg, doc)

        if idx < len(report_rows) - 1:
            doc.add_paragraph()

    doc.add_paragraph()
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run(
        f"Work Log Report  –  {week_str}  –  {datetime.date.today().strftime('%Y-%m-%d')}"
    ).font.size = Pt(8)

    doc.save(out_buf)


HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
<meta name="apple-mobile-web-app-title" content="Work Log">
<meta name="theme-color" content="#0f172a">
<title>Work Log Journal</title>
<link rel="manifest" href="/manifest.json">
<link rel="apple-touch-icon" href="/apple-touch-icon.png">
<link rel="apple-touch-icon" sizes="180x180" href="/apple-touch-icon.png">
<style>
  @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;600&display=swap');
  :root {
    --bg:       #0f172a;
    --bg2:      #1e293b;
    --bg3:      #334155;
    --border:   #475569;
    --accent:   #38bdf8;
    --accent2:  #0ea5e9;
    --fg:       #f1f5f9;
    --fg2:      #94a3b8;
    --green:    #4ade80;
    --yellow:   #fbbf24;
    --red:      #f87171;
    --grey:     #64748b;
    --radius:   10px;
    --font:     'DM Sans', system-ui, -apple-system, sans-serif;
    --mono:     'JetBrains Mono', monospace;
    --emoji:    'Noto Color Emoji', 'Segoe UI Emoji', 'Apple Color Emoji', 'Segoe UI Symbol';
  }
  [data-theme="light"] {
    --bg:       #f8fafc;
    --bg2:      #ffffff;
    --bg3:      #e2e8f0;
    --border:   #cbd5e1;
    --accent:   #0ea5e9;
    --accent2:  #0284c7;
    --fg:       #0f172a;
    --fg2:      #64748b;
    --green:    #16a34a;
    --yellow:   #d97706;
    --red:      #dc2626;
    --grey:     #9ca3af;
  }
  * { box-sizing:border-box; margin:0; padding:0; }
  html,body { height:100%; font-family:var(--font); background:var(--bg); color:var(--fg); }
  .icon { font-family:var(--emoji), var(--font); font-style:normal; }

  /* ── Layout ── */
  #app { display:flex; flex-direction:column; height:100vh; overflow:hidden; }
  #toolbar { display:flex; align-items:center; gap:8px; padding:10px 14px;
             background:var(--bg2); border-bottom:1px solid var(--border);
             flex-shrink:0; flex-wrap:wrap; }
  #toolbar h1 { font-size:16px; font-weight:700; color:var(--accent);
                letter-spacing:-.3px; margin-right:8px; white-space:nowrap; }
  #db-badge { font-size:11px; color:var(--fg2); background:var(--bg3);
              padding:2px 8px; border-radius:20px; white-space:nowrap; }
  #version-badge { font-size:10px; color:var(--fg2); background:var(--bg3);
              padding:2px 8px; border-radius:20px; white-space:nowrap; }
  .spacer { flex:1; }
  #filter-bar { display:flex; align-items:center; gap:8px; padding:8px 14px;
                background:var(--bg); border-bottom:1px solid var(--border);
                flex-shrink:0; flex-wrap:wrap; }
  #filter-bar label { font-size:12px; color:var(--fg2); }
  #main { flex:1; overflow:hidden; display:flex; flex-direction:column; }

  /* ── Buttons ── */
  .btn { display:inline-flex; align-items:center; gap:5px; padding:6px 12px;
         border:none; border-radius:6px; cursor:pointer; font-size:12px;
         font-weight:600; font-family:var(--font); transition:.15s; white-space:nowrap; }
  .btn:active { opacity:.8; transform:scale(.97); }
  .btn-accent  { background:var(--accent2); color:#fff; }
  .btn-green   { background:#16a34a; color:#fff; }
  .btn-red     { background:#dc2626; color:#fff; }
  .btn-yellow  { background:#d97706; color:#fff; }
  .btn-ghost   { background:var(--bg3); color:var(--fg); }
  .btn-word    { background:#2563eb; color:#fff; }
  .btn-sm      { padding:4px 9px; font-size:11px; }

  /* ── Inputs ── */
  input, select, textarea {
    background:var(--bg2); color:var(--fg); border:1px solid var(--border);
    border-radius:6px; padding:6px 10px; font-size:13px; font-family:var(--font);
    width:100%; outline:none; transition:border .15s; }
  input:focus, select:focus, textarea:focus { border-color:var(--accent); }
  input[type="date"] { width:140px; }
  select { width:auto; min-width:120px; }
  textarea { resize:vertical; font-family:var(--mono); font-size:12px; }

  /* ── Table ── */
  #table-wrap { flex:1; overflow:auto; }
  table { width:100%; border-collapse:collapse; font-size:12px; }
  thead th { position:sticky; top:0; background:var(--bg2); padding:9px 10px;
             text-align:left; font-size:11px; font-weight:600; color:var(--fg2);
             border-bottom:1px solid var(--border); white-space:nowrap;
             cursor:pointer; user-select:none; }
  thead th:hover { color:var(--accent); }
  tbody tr { border-bottom:1px solid var(--border); cursor:pointer; transition:.1s; }
  tbody tr:hover { background:var(--bg3); }
  tbody tr.selected { background:color-mix(in srgb, var(--accent) 15%, var(--bg)); }
  tbody td { padding:7px 10px; vertical-align:top; max-width:260px;
             overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }
  tbody td.wl-cell { white-space:normal; max-height:60px; overflow:hidden; max-width:280px !important; width:280px !important; } /* v1.0.5: Worklogs wider */
  tbody td.mchp-device-cell { max-width:150px; white-space:normal; word-wrap:break-word; }

  /* v1.0.5: Column width adjustments - forced with !important */
  thead th:nth-child(8), tbody td:nth-child(8) { width:280px !important; max-width:280px !important; min-width:280px !important; } /* Worklogs - wider */
  thead th:nth-child(9), tbody td:nth-child(9) { width:140px !important; max-width:140px !important; } /* Milestone - narrower */
  thead th:nth-child(10), tbody td:nth-child(10) { width:140px !important; max-width:140px !important; } /* To-Do - narrower */
  .badge { display:inline-block; padding:2px 8px; border-radius:20px;
           font-size:10px; font-weight:600; }
  .badge-done   { background:color-mix(in srgb,var(--green) 20%,transparent);  color:var(--green); }
  .badge-wip    { background:color-mix(in srgb,var(--yellow) 20%,transparent); color:var(--yellow); }
  .badge-cancel { background:color-mix(in srgb,var(--red) 20%,transparent);    color:var(--red); }
  .badge-other  { background:var(--bg3); color:var(--fg2); }
  tr.archived td { opacity:.5; text-decoration:line-through; color:var(--fg2); }

  /* ── Modal ── */
  #modal-overlay { display:none; position:fixed; inset:0; background:rgba(0,0,0,.65);
                   z-index:100; align-items:center; justify-content:center; padding:12px; }
  #modal-overlay.open { display:flex; }
  #modal { background:var(--bg2); border-radius:14px; width:100%; max-width:860px;
           max-height:95vh; display:flex; flex-direction:column;
           border:1px solid var(--border); box-shadow:0 24px 80px rgba(0,0,0,.5); }
  #modal-header { display:flex; align-items:center; justify-content:space-between;
                  padding:14px 18px; border-bottom:1px solid var(--border); flex-shrink:0; }
  #modal-header h2 { font-size:15px; font-weight:700; }
  #modal-body { overflow-y:auto; padding:16px 18px; flex:1; }
  #modal-footer { padding:12px 18px; border-top:1px solid var(--border);
                  display:flex; justify-content:flex-end; gap:8px; flex-shrink:0; }
  .form-grid { display:grid; grid-template-columns:1fr 1fr; gap:10px 18px; }
  .form-row { display:flex; flex-direction:column; gap:4px; }
  .form-row label { font-size:11px; font-weight:600; color:var(--fg2); text-transform:uppercase; letter-spacing:.4px; }
  .form-full { grid-column:1/-1; }
  .char-hint { font-size:10px; color:var(--fg2); text-align:right; margin-top:2px; }

  /* ── Worklogs editor ── */
  #wl-tabs, #ps-tabs, #todo-tabs { display:flex; gap:4px; margin-bottom:6px; }
  #wl-edit, #ps-edit, #todo-edit  { display:block; }
  #wl-preview, #ps-preview, #todo-preview {
    display:none; background:var(--bg); border:1px solid var(--border);
    border-radius:6px; padding:12px; font-size:13px; line-height:1.7; }
  #wl-preview { min-height:200px; }
  #ps-preview { min-height:100px; }
  #todo-preview { min-height:60px; }
  #wl-preview h1, #ps-preview h1, #todo-preview h1 { font-size:1.4em; color:var(--accent); margin:.4em 0; }
  #wl-preview h2, #ps-preview h2, #todo-preview h2 { font-size:1.2em; color:var(--accent); margin:.3em 0; }
  #wl-preview h3, #ps-preview h3, #todo-preview h3 { font-size:1em;   color:var(--accent); margin:.2em 0; }
  #wl-preview strong, #ps-preview strong, #todo-preview strong { font-weight:700; }
  #wl-preview em, #ps-preview em, #todo-preview em { font-style:italic; }
  #wl-preview code, #ps-preview code, #todo-preview code {
    background:var(--bg3); color:#f472b6; padding:1px 5px;
    border-radius:4px; font-family:var(--mono); font-size:.9em; }
  #wl-preview .stamp { color:var(--accent); font-weight:700; font-size:.85em; }
  #wl-preview ul, #wl-preview ol, #ps-preview ul, #ps-preview ol,
  #todo-preview ul, #todo-preview ol { padding-left:1.5em; margin:.2em 0; }
  #wl-preview ul li, #ps-preview ul li, #todo-preview ul li { list-style-type: disc; }
  #wl-preview ul ul li, #ps-preview ul ul li, #todo-preview ul ul li { list-style-type: circle; }
  #wl-preview ul ul ul li, #ps-preview ul ul ul li, #todo-preview ul ul ul li { list-style-type: square; }
  #wl-preview ol li, #ps-preview ol li, #todo-preview ol li { list-style-type: decimal; }
  #wl-preview ol ol li, #ps-preview ol ol li, #todo-preview ol ol li { list-style-type: lower-alpha; }
  #wl-preview ol ol ol li, #ps-preview ol ol ol li, #todo-preview ol ol ol li { list-style-type: lower-roman; }
  #wl-preview li, #ps-preview li, #todo-preview li { margin:.1em 0; }
  #wl-preview hr, #ps-preview hr, #todo-preview hr { border:none; border-top:1px solid var(--border); margin:.8em 0; }
  #wl-preview img { max-width:100%; border-radius:6px; margin:.5em 0; }
  .md-hint { font-size:10.5px; color:var(--fg2); background:var(--bg3);
             padding:5px 10px; border-radius:6px; margin-top:5px; line-height:1.8; }
  .paste-zone { border:2px dashed var(--border); border-radius:8px; padding:10px;
                text-align:center; font-size:12px; color:var(--fg2); cursor:pointer;
                margin-top:6px; transition:.15s; }
  .paste-zone:hover, .paste-zone.drag-over { border-color:var(--accent); color:var(--accent); }
  .img-thumb { max-width:160px; max-height:120px; border-radius:6px;
               margin:4px; vertical-align:top; cursor:zoom-in;
               transition:opacity .15s; }
  .img-thumb:hover { opacity:.85; }
  .file-attach-link { display:inline-flex;align-items:center;gap:4px;
    padding:2px 8px;border-radius:4px;background:var(--bg3);
    color:var(--accent);text-decoration:none;font-size:0.9em;
    border:1px solid var(--border); }
  .file-attach-link:hover { opacity:.8; }

  /* Lightbox */
  #lightbox { display:none; position:fixed; inset:0; z-index:999;
              background:rgba(0,0,0,.88); align-items:center;
              justify-content:center; cursor:zoom-out; }
  #lightbox.open { display:flex; }
  #lightbox-container { position:relative; width:600px; height:450px;
                        min-width:200px; min-height:150px;
                        max-width:95vw; max-height:95vh;
                        resize:both; overflow:auto;
                        border-radius:8px; box-shadow:0 8px 48px rgba(0,0,0,.7); }
  #lightbox-img { width:100%; height:100%; border-radius:8px;
                  object-fit:contain; cursor:default;
                  display:block; }
  #lightbox-close { position:fixed; top:16px; right:20px; background:rgba(255,255,255,.15);
                    border:none; color:#fff; font-size:22px; width:36px; height:36px;
                    border-radius:50%; cursor:pointer; line-height:1;
                    display:flex; align-items:center; justify-content:center; }
  #lightbox-close:hover { background:rgba(255,255,255,.3); }

  /* ── Import dialog ── */
  #import-result { margin-top:10px; max-height:220px; overflow-y:auto; }
  .ir-row { display:flex; gap:8px; padding:5px 8px; border-radius:6px;
            font-size:12px; margin-bottom:3px; }
  .ir-added   { background:color-mix(in srgb,var(--green) 12%,transparent); }
  .ir-merged  { background:color-mix(in srgb,var(--yellow)12%,transparent); }
  .ir-skipped { background:var(--bg3); }

  /* ── Toast ── */
  #toast { position:fixed; bottom:20px; left:50%; transform:translateX(-50%) translateY(40px);
           background:var(--accent2); color:#fff; padding:9px 18px; border-radius:8px;
           font-size:13px; font-weight:500; opacity:0; transition:.3s; pointer-events:none;
           z-index:999; white-space:nowrap; }
  #toast.show { opacity:1; transform:translateX(-50%) translateY(0); }

  /* ── Theme toggle ── */
  .theme-btn { background:none; border:none; cursor:pointer; font-size:16px;
               padding:4px; border-radius:6px; transition:.15s; }
  .theme-btn:hover { background:var(--bg3); }

  /* ── Sync status bar ── */
  #sync-bar { display:flex; align-items:center; gap:8px; padding:4px 14px;
              background:var(--bg2); border-bottom:1px solid var(--border);
              font-size:11px; color:var(--fg2); flex-shrink:0; flex-wrap:wrap; }
  .sync-dot { width:8px; height:8px; border-radius:50%; flex-shrink:0; }
  .sync-online  { background:#22c55e; }
  .sync-offline { background:#f59e0b; animation:pulse 2s infinite; }
  .sync-error   { background:#ef4444; }
  @keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.4} }
  .conflict-badge { background:#ef4444; color:#fff; border-radius:10px;
                    padding:1px 7px; font-size:10px; font-weight:700;
                    cursor:pointer; margin-left:4px; }

  /* ── Conflict modal ── */
  #conflict-overlay { display:none; position:fixed; inset:0; background:rgba(0,0,0,.7);
                      z-index:150; align-items:center; justify-content:center; padding:16px; }
  #conflict-overlay.open { display:flex; }
  #conflict-modal { background:var(--bg2); border-radius:14px; width:100%;
                    max-width:820px; max-height:90vh; display:flex;
                    flex-direction:column; border:1px solid var(--border); }
  .conflict-card { border:1px solid var(--border); border-radius:8px;
                   padding:12px; margin-bottom:10px; }
  .conflict-col { flex:1; min-width:0; }
  .conflict-col h4 { font-size:11px; font-weight:700; margin-bottom:6px;
                     text-transform:uppercase; letter-spacing:.5px; }
  .local-col  h4 { color:#f59e0b; }
  .cloud-col  h4 { color:#38bdf8; }
  .conf-field { font-size:11px; margin-bottom:3px; }
  .conf-field strong { color:var(--fg2); min-width:90px; display:inline-block; }
  .conf-diff  { background:color-mix(in srgb, #ef4444 12%, transparent);
                border-radius:4px; padding:1px 4px; }

  @media (max-width:600px) {
    .form-grid { grid-template-columns:1fr; }
    .form-full { grid-column:1; }
    thead th:nth-child(n+7) { display:none; }
    tbody td:nth-child(n+7) { display:none; }
    #toolbar { gap:5px; }
    .btn span { display:none; }
  }
</style>
</head>
<body data-theme="dark">
<div id="app">

  <!-- Toolbar -->
  <div id="toolbar">
    <h1><i class="icon">📋</i> Work Log</h1>
    <span id="db-badge">loading…</span>
    <button class="btn btn-green"  onclick="openNew()"><i class="icon">➕</i> <span>New</span></button>
    <button class="btn btn-yellow" onclick="openEdit()"><i class="icon">✏️</i> <span>Edit</span></button>
    <button class="btn btn-red"    onclick="deleteRecord()"><i class="icon">🗑</i> <span>Delete</span></button>
    <button class="btn btn-ghost"  onclick="openImport()"><i class="icon">📥</i> <span>Import</span></button>
    <button class="btn btn-word"   onclick="exportWord()"><i class="icon">📄</i> <span>Word</span></button>
    <button class="btn btn-ghost"  onclick="exportExcel()"><i class="icon">📊</i> <span>Excel</span></button>
    <button class="btn btn-ghost"  onclick="exportPdf()"><i class="icon">📑</i> <span>PDF</span></button>
    <div class="spacer"></div>
    <span id="db-size-badge" style="font-size:10px;color:var(--fg2);margin-right:8px">DB: --</span>
    <span id="egress-badge" style="font-size:10px;color:var(--fg2);margin-right:8px;display:none">Egress: --</span>
    <span id="sync-status-badge" style="font-size:10px;color:var(--fg2);margin-right:8px" title="Sync status">⏳ Checking...</span>
    <span id="version-badge">v{{APP_VERSION}} by Keynes Hsu</span>
    <button class="theme-btn" onclick="cycleTheme()" title="Toggle theme" id="theme-icon"><i class="icon">🌙</i></button>
    <button class="btn btn-ghost btn-sm" onclick="refreshList()" title="Refresh"><i class="icon">🔄</i></button>
  </div>

  <!-- Filter bar -->
  <div id="filter-bar">
    <label>Customer</label>
    <input id="f-cust" style="width:130px" placeholder="filter…" oninput="refreshList()">
    <label>Project</label>
    <input id="f-proj" style="width:130px" placeholder="filter…" oninput="refreshList()">
    <label>Status</label>
    <select id="f-status" onchange="refreshList()"><option value="">All</option></select>
    <label>Category</label>
    <select id="f-cat" onchange="refreshList()"><option value="">All</option></select>
    <label>Archive</label>
    <select id="f-arch" onchange="refreshList()">
      <option value="">All</option><option>No</option><option>Yes</option>
    </select>
    <button class="btn btn-ghost btn-sm" onclick="clearFilters()">Clear</button>
  </div>

  <!-- Sync status bar -->
  <div id="sync-bar">
    <span class="sync-dot sync-offline" id="sync-dot"></span>
    <span id="sync-text">Checking…</span>
    <span id="sync-pending" style="color:var(--yellow)"></span>
    <span id="conflict-badge-wrap"></span>
    <div class="spacer"></div>
    <button class="btn btn-ghost btn-sm" onclick="syncNow()" id="sync-btn"
            title="Sync now">↻ Sync</button>
    <button class="btn btn-ghost btn-sm" onclick="openSyncSettings()" title="Sync settings">⚙ DB</button>
  </div>

  <!-- Table -->
  <div id="main">
    <div id="table-wrap">
      <table id="records-table">
        <thead>
          <tr>
            <th style="width:36px;text-align:center">
              <input type="checkbox" id="chk-all" title="Select all"
                     onchange="toggleAll(this.checked)" style="width:auto;cursor:pointer">
            </th>
            <th onclick="sortBy('week')">Week ⇅</th>
            <th onclick="sortBy('category')">Category ⇅</th>
            <th onclick="sortBy('customer')">Customer ⇅</th>
            <th onclick="sortBy('project_name')">Project Name ⇅</th>
            <th onclick="sortBy('status')">Status ⇅</th>
            <th onclick="sortBy('task_summary')">Task Summary ⇅</th>
            <th>Worklogs</th>
            <th onclick="sortBy('project_schedule')">Milestone ⇅</th>
            <th onclick="sortBy('todo_content')">To-Do ⇅</th>
            <th onclick="sortBy('application')">Application ⇅</th>
            <th onclick="sortBy('sso_modeln')">SSO/ModelN# ⇅</th>
            <th onclick="sortBy('mchp_device')">MCHP Device ⇅</th>
            <th onclick="sortBy('ear')">EAR(K$) ⇅</th>
            <th onclick="sortBy('bu')">BU ⇅</th>
            <th onclick="sortBy('last_update')">Last Update ⇅</th>
            <th onclick="sortBy('create_date')">Create Date ⇅</th>
            <th onclick="sortBy('due_date')">Due Date ⇅</th>
            <th onclick="sortBy('archive')">Archive ⇅</th>
          </tr>
        </thead>
        <tbody id="records-body"></tbody>
      </table>
    </div>
  </div>
</div>

<!-- Record Modal -->
<div id="modal-overlay">
  <div id="modal">
    <div id="modal-header">
      <h2 id="modal-title">New Record</h2>
      <span id="auto-save-indicator" style="font-size:11px;color:var(--green);opacity:0;transition:opacity 0.3s;margin-left:12px"></span>
      <button class="btn btn-ghost btn-sm" onclick="closeModal()">✕</button>
    </div>
    <div id="modal-body">
      <div class="form-grid">
        <!-- Col 1 -->
        <div class="form-row">
          <label>Create Date</label>
          <input type="date" id="f-create_date">
        </div>
        <div class="form-row">
          <label>Customer <small style="color:var(--fg2)">(max 20)</small></label>
          <input id="f-customer" maxlength="20" oninput="updateCharHint(this,'ch-cust')">
          <div class="char-hint" id="ch-cust">0/20</div>
        </div>
        <div class="form-row">
          <label>Due Date</label>
          <input type="date" id="f-due_date">
        </div>
        <div class="form-row">
          <label>Project Name <small style="color:var(--fg2)">(max 20)</small></label>
          <input id="f-project_name" maxlength="20" oninput="updateCharHint(this,'ch-proj')">
          <div class="char-hint" id="ch-proj">0/20</div>
        </div>
        <div class="form-row">
          <label>SSO/ModelN# <small style="color:var(--fg2)">(max 10)</small></label>
          <input id="f-sso_modeln" maxlength="10" oninput="updateCharHint(this,'ch-sso')">
          <div class="char-hint" id="ch-sso">0/10</div>
        </div>
        <div class="form-row">
          <label>EAR(K$) <small style="color:var(--fg2)">(max 20)</small></label>
          <input id="f-ear" maxlength="20" oninput="updateCharHint(this,'ch-ear')">
          <div class="char-hint" id="ch-ear">0/20</div>
        </div>
        <div class="form-row">
          <label>Application <small style="color:var(--fg2)">(max 50)</small></label>
          <input id="f-application" maxlength="50" oninput="updateCharHint(this,'ch-app')">
          <div class="char-hint" id="ch-app">0/50</div>
        </div>
        <div class="form-row">
          <label>Last Update</label>
          <input type="date" id="f-last_update" oninput="autoWeek()">
        </div>
        <div class="form-row">
          <label>MCHP Device</label>
          <input id="f-mchp_device">
        </div>
        <div class="form-row">
          <label>Week <small style="color:var(--fg2)">(auto)</small></label>
          <input id="f-week" readonly style="background:var(--bg3);width:80px">
        </div>
        <div class="form-row">
          <label>BU</label>
          <select id="f-bu">
            <option value="">--</option>
            <option value="DCS">DCS</option>
            <option value="NCS">NCS</option>
          </select>
        </div>
        <div class="form-row">
          <label>Category</label>
          <select id="f-category"></select>
        </div>
        <div class="form-row">
          <label>Status</label>
          <select id="f-status-field"></select>
        </div>
        <div class="form-row" style="align-self:end">
          <label><input type="checkbox" id="f-archive" style="width:auto;margin-right:6px">Archive</label>
        </div>
        <!-- Full width -->
        <div class="form-row form-full">
          <label>Task Summary</label>
          <textarea id="f-task_summary" rows="3"></textarea>
        </div>
        <!-- Worklogs -->
        <div class="form-row form-full">
          <div style="display:flex;align-items:center;gap:8px;margin-bottom:6px">
            <label style="margin:0">Worklogs</label>
            <div id="wl-tabs">
              <button class="btn btn-accent btn-sm" onclick="setWlTab('edit')"><i class="icon">✏️</i> Edit</button>
              <button class="btn btn-ghost btn-sm" onclick="setWlTab('preview')">👁 Preview</button>
              <button class="btn btn-ghost btn-sm" onclick="insertStamp()"><i class="icon">➕</i> Add Entry</button>
            </div>
          </div>
          <div id="wl-edit">
            <textarea id="f-worklogs" rows="12" placeholder="── 2025-03-01 09:00 ──&#10;Write your worklog here…&#10;Markdown supported."></textarea>
            <div style="display:flex;gap:6px;margin-top:4px;">
              <div class="paste-zone" id="paste-zone" style="flex:1"
                   onclick="document.getElementById('img-upload').click()"
                   ondragover="event.preventDefault();this.classList.add('drag-over')"
                   ondragleave="this.classList.remove('drag-over')"
                   ondrop="handleDrop(event)">
                🖼 Image — Click, drag & drop, or Ctrl+V
              </div>
              <div class="paste-zone" id="file-zone" style="flex:1"
                   onclick="document.getElementById('file-upload').click()"
                   ondragover="event.preventDefault();this.classList.add('drag-over')"
                   ondragleave="this.classList.remove('drag-over')"
                   ondrop="handleFileDrop(event)">
                📎 File — Click or drag & drop (max 5 MB)
              </div>
            </div>
            <input type="file" id="img-upload" accept="image/*" style="display:none" onchange="handleFileSelect(event)">
            <input type="file" id="file-upload" style="display:none" onchange="handleNonImageFileSelect(event)">
            <div id="img-preview-area"></div>
          </div>
          <div id="wl-preview" style="display:none"></div>
          <div class="md-hint">
            <strong>Markdown:</strong>
            <code># H1</code> <code>## H2</code> <code>### H3</code> &nbsp;|&nbsp;
            <code>**bold**</code> <code>*italic*</code> <code>__underline__</code> <code>~~strikethrough~~</code> <code>`code`</code> &nbsp;|&nbsp;
            <code>- bullet</code> <code>1. list</code>
            <code>&nbsp;&nbsp;- sub</code> <code>&nbsp;&nbsp;&nbsp;&nbsp;- sub-sub</code>
            (indent 2 spaces = next level) &nbsp;|&nbsp;
            <code>---</code> &nbsp;|&nbsp;
            Ctrl+V / drag &amp; drop image
          </div>
        </div>
        <!-- Milestone -->
        <div class="form-row form-full">
          <div style="display:flex;align-items:center;gap:8px;margin-bottom:6px">
            <label style="margin:0">Milestone</label>
            <div id="ps-tabs">
              <button class="btn btn-accent btn-sm" onclick="setPsTab('edit')"><i class="icon">✏️</i> Edit</button>
              <button class="btn btn-ghost btn-sm" onclick="setPsTab('preview')">👁 Preview</button>
            </div>
          </div>
          <div id="ps-edit">
            <textarea id="f-project_schedule" rows="6" placeholder="Project milestones and timeline...&#10;Markdown supported."></textarea>
          </div>
          <div id="ps-preview"></div>
          <div class="md-hint">
            <strong>Markdown:</strong>
            <code>**bold**</code> <code>*italic*</code> <code>__underline__</code> <code>~~strikethrough~~</code> <code>`code`</code> &nbsp;|&nbsp;
            <code>- bullet</code> <code>1. list</code>
            <code>&nbsp;&nbsp;- sub</code> (indent 2 spaces = next level)
          </div>
        </div>
        <!-- To-Do -->
        <div class="form-row form-full">
          <div style="display:flex;align-items:center;gap:8px;margin-bottom:6px">
            <label style="margin:0">To-Do</label>
            <div id="todo-tabs">
              <button class="btn btn-accent btn-sm" onclick="setTodoTab('edit')"><i class="icon">✏️</i> Edit</button>
              <button class="btn btn-ghost btn-sm" onclick="setTodoTab('preview')">👁 Preview</button>
            </div>
          </div>
          <div id="todo-edit">
            <textarea id="f-todo_content" rows="3" placeholder="What needs to be done...&#10;Markdown supported."></textarea>
          </div>
          <div id="todo-preview"></div>
        </div>
        <div class="form-row">
          <label>Due Date</label>
          <input type="date" id="f-todo_due_date">
        </div>
      </div>
    </div>
    <div id="modal-footer">
      <button class="btn btn-ghost" onclick="closeModal()">Cancel</button>
      <button class="btn btn-accent" onclick="saveRecord()"><i class="icon">💾</i> Save</button>
    </div>
  </div>
</div>

<!-- Import Modal -->
<div id="import-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.65);z-index:110;align-items:center;justify-content:center;padding:12px;flex-direction:column">
  <div style="background:var(--bg2);border-radius:14px;padding:20px;max-width:520px;width:100%;border:1px solid var(--border)">
    <h2 style="margin-bottom:12px;font-size:15px"><i class="icon">📥</i> 資料庫匯入</h2>
    <p style="font-size:12px;color:var(--fg2);margin-bottom:10px">
      上傳 SQLite 資料庫檔案（.db）並匯入所有記錄到本機。相同 Customer + Project 的記錄會合併 Worklogs。
    </p>
    <div style="margin-bottom:12px">
      <label style="display:block;font-size:12px;margin-bottom:6px;color:var(--fg1)">選擇資料庫檔案</label>
      <input type="file" id="import-db-file" accept=".db,.sqlite,.sqlite3"
             style="width:100%;padding:8px;border:1px solid var(--border);border-radius:6px;background:var(--bg1);color:var(--fg1);font-size:12px">
    </div>
    <button class="btn btn-accent" onclick="doImport()">開始匯入</button>
    <button class="btn btn-ghost" onclick="closeImport()" style="margin-left:8px">取消</button>
    <div id="import-result" style="margin-top:12px"></div>
  </div>
</div>

<!-- Sync Settings Overlay -->
<div id="sync-settings-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.65);z-index:120;align-items:center;justify-content:center;padding:12px">
  <div style="background:var(--bg2);border-radius:14px;padding:24px;max-width:460px;width:100%;border:1px solid var(--border)">
    <h2 style="margin-bottom:4px;font-size:15px">⚙ Sync Settings</h2>
    <p style="font-size:12px;color:var(--fg2);margin-bottom:16px">設定雲端資料庫連線。變更會立即套用並儲存至 .env。</p>

    <div style="margin-bottom:12px">
      <label style="display:block;font-size:12px;font-weight:600;margin-bottom:4px">資料庫類型</label>
      <select id="ss-mode" onchange="onSsModeChange()" style="width:100%;padding:6px 8px;border-radius:6px;border:1px solid var(--border);background:var(--bg);color:var(--fg);font-size:13px">
        <option value="local">SQLite（本機）</option>
        <option value="pocketbase">PocketBase</option>
        <option value="postgres">PostgreSQL</option>
      </select>
    </div>

    <div id="ss-pb-section">
      <div style="margin-bottom:10px">
        <label style="display:block;font-size:12px;font-weight:600;margin-bottom:4px">PocketBase URL</label>
        <input id="ss-pb-url" type="text" placeholder="http://127.0.0.1:8090"
               style="width:100%;box-sizing:border-box;padding:6px 8px;border-radius:6px;border:1px solid var(--border);background:var(--bg);color:var(--fg);font-size:13px">
        <div style="font-size:11px;color:var(--fg2);margin-top:3px">e.g. http://127.0.0.1:8090 或 https://your-pb-server.com</div>
      </div>
      <div style="margin-bottom:10px">
        <label style="display:block;font-size:12px;font-weight:600;margin-bottom:4px">PocketBase Token（選填）</label>
        <input id="ss-pb-token" type="password" placeholder="留空表示不變更"
               style="width:100%;box-sizing:border-box;padding:6px 8px;border-radius:6px;border:1px solid var(--border);background:var(--bg);color:var(--fg);font-size:13px">
      </div>
    </div>

    <div id="ss-pg-section" style="display:none">
      <div style="margin-bottom:10px">
        <label style="display:block;font-size:12px;font-weight:600;margin-bottom:4px">PostgreSQL URL</label>
        <input id="ss-pg-url" type="text" placeholder="postgresql://user:pass@host:5432/db"
               style="width:100%;box-sizing:border-box;padding:6px 8px;border-radius:6px;border:1px solid var(--border);background:var(--bg);color:var(--fg);font-size:13px">
        <div style="font-size:11px;color:var(--fg2);margin-top:3px">格式：postgresql://使用者:密碼@主機IP:5432/資料庫名</div>
      </div>
    </div>

    <div style="display:flex;gap:8px;justify-content:flex-end;margin-top:16px">
      <button class="btn btn-ghost btn-sm" onclick="closeSyncSettings()">Cancel</button>
      <button class="btn btn-accent btn-sm" onclick="saveSyncSettings()"><i class="icon">💾</i> Save &amp; Apply</button>
    </div>
  </div>
</div>

<!-- Week Selection Modal for Word Export -->
<div id="week-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.65);z-index:115;align-items:center;justify-content:center;padding:12px">
  <div style="background:var(--bg2);border-radius:14px;padding:20px;max-width:400px;width:100%;border:1px solid var(--border)">
    <h2 style="margin-bottom:12px;font-size:15px"><i class="icon">📄</i> Export to Word</h2>
    <p style="font-size:12px;color:var(--fg2);margin-bottom:12px">
      Select which week's worklogs to include in the report:
    </p>
    <div style="margin-bottom:14px">
      <select id="week-select" style="width:100%;padding:10px;font-size:13px">
        <option value="all">All Worklogs</option>
      </select>
    </div>
    <div id="week-info" style="font-size:11px;color:var(--fg2);margin-bottom:14px;padding:8px;background:var(--bg3);border-radius:6px;display:none">
    </div>
    <div style="display:flex;gap:8px;justify-content:flex-end">
      <button class="btn btn-ghost" onclick="closeWeekSelect()">Cancel</button>
      <button class="btn btn-word" onclick="confirmWeekExport()"><i class="icon">📄</i> Export</button>
    </div>
  </div>
</div>

<!-- Excel Export Options Modal -->
<div id="excel-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.65);z-index:115;align-items:center;justify-content:center;padding:12px">
  <div style="background:var(--bg2);border-radius:14px;padding:20px;max-width:420px;width:100%;border:1px solid var(--border)">
    <h2 style="margin-bottom:12px;font-size:15px"><i class="icon">📊</i> Export to Excel</h2>
    <p style="font-size:12px;color:var(--fg2);margin-bottom:14px">
      Select export format:
    </p>

    <!-- Format Selection -->
    <div style="margin-bottom:16px">
      <label style="display:block;margin-bottom:8px;cursor:pointer;padding:10px;background:var(--bg3);border-radius:8px;border:2px solid transparent" id="format-wlr-label">
        <input type="radio" name="excel-format" value="WLR" id="format-wlr" checked style="margin-right:8px">
        <strong>WLR Format</strong> — Work Log Report
        <div style="font-size:11px;color:var(--fg2);margin-left:24px;margin-top:4px">
          Comprehensive reporting format with all key fields
        </div>
      </label>
      <label style="display:block;cursor:pointer;padding:10px;background:var(--bg3);border-radius:8px;border:2px solid transparent" id="format-todo-label">
        <input type="radio" name="excel-format" value="TODO" id="format-todo" style="margin-right:8px">
        <strong>To-Do Format</strong> — Task List
        <div style="font-size:11px;color:var(--fg2);margin-left:24px;margin-top:4px">
          Simple task list with due dates and checkboxes
        </div>
      </label>
    </div>

    <!-- Date Filter (shown only for To-Do format) -->
    <div id="date-filter-section" style="display:none;margin-bottom:16px;padding:12px;background:var(--bg3);border-radius:8px">
      <p style="font-size:12px;font-weight:600;margin-bottom:10px">Date Filter:</p>
      <label style="display:block;margin-bottom:6px;cursor:pointer">
        <input type="radio" name="date-filter" value="all" id="filter-all" checked style="margin-right:8px">
        All Due Dates
      </label>
      <label style="display:block;cursor:pointer">
        <input type="radio" name="date-filter" value="week_plus_1" id="filter-week" style="margin-right:8px">
        Week+1 (Next 7 days only)
      </label>
    </div>

    <div style="display:flex;gap:8px;justify-content:flex-end">
      <button class="btn btn-ghost" onclick="closeExcelSelect()">Cancel</button>
      <button class="btn btn-accent" onclick="confirmExcelExport()"><i class="icon">📊</i> Export</button>
    </div>
  </div>
</div>

<div id="stamp-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.65);
     z-index:200;align-items:center;justify-content:center">
  <div style="background:var(--bg2);border-radius:12px;padding:20px;
              min-width:280px;border:1px solid var(--border);box-shadow:0 16px 48px rgba(0,0,0,.5)">
    <h3 style="margin-bottom:14px;font-size:14px">📅  Insert Worklog Entry</h3>
    <label style="font-size:11px;font-weight:600;color:var(--fg2);text-transform:uppercase;
                  letter-spacing:.4px;display:block;margin-bottom:4px">Date</label>
    <input type="date" id="stamp-date" style="width:100%;margin-bottom:14px">
    <div style="display:flex;gap:8px;justify-content:flex-end">
      <button class="btn btn-ghost btn-sm" onclick="closeStamp()">Cancel</button>
      <button class="btn btn-accent btn-sm" onclick="confirmStamp()">Insert</button>
    </div>
  </div>
</div>

<!-- Conflict Resolution Modal -->
<div id="conflict-overlay">
  <div id="conflict-modal">
    <div style="display:flex;align-items:center;justify-content:space-between;
                padding:14px 18px;border-bottom:1px solid var(--border);
                background:#7c1d1d;border-radius:14px 14px 0 0">
      <span style="font-weight:700;color:#fff;font-size:15px">
        ⚠️  Sync Conflicts — Please Review
      </span>
      <button class="btn btn-ghost btn-sm" onclick="closeConflicts()"
              style="color:#fff;background:rgba(255,255,255,.15)">✕</button>
    </div>
    <div style="padding:16px 18px;overflow-y:auto;flex:1">
      <p style="font-size:12px;color:var(--fg2);margin-bottom:14px">
        These records were edited locally while offline AND also updated in the cloud.
        Choose which version to keep for each conflict.
      </p>
      <div id="conflict-list"></div>
    </div>
    <div style="padding:12px 18px;border-top:1px solid var(--border);
                display:flex;justify-content:flex-end;gap:8px">
      <button class="btn btn-ghost" onclick="closeConflicts()">Close</button>
    </div>
  </div>
</div>

<div id="lightbox" onclick="closeLightbox(event)">
  <button id="lightbox-close" onclick="closeLightbox()">✕</button>
  <div id="lightbox-container" onclick="event.stopPropagation()">
    <img id="lightbox-img" src="" alt="preview">
  </div>
</div>

<div id="toast"></div>

<script>
// ── State ─────────────────────────────────────────────────────────────────────
let records = [], selectedId = null, editingId = null;
let sortCol = 'last_update', sortAsc = false;
let cfg = {};
let themes = ['dark','light','system'];
let themeIdx = 0;
let pendingImages = []; // [{uid, b64}]
let pendingFiles   = []; // [{uid, name, mimeType, b64, size}]
let checkedIds = new Set(); // IDs selected for Word export

// v1.0.5: Update sync status badge in footer
function updateSyncStatusBadge(status, data = null) {
  const el = document.getElementById('sync-status-badge');
  if (!el) return;

  const states = {
    'synced': { text: '✅ Synced', color: 'var(--green)', title: 'All changes synced' },
    'syncing': { text: '🔄 Syncing...', color: 'var(--blue)', title: 'Synchronizing with server' },
    'pending': { text: '⏳ Pending', color: 'var(--orange)', title: 'Changes waiting to sync' },
    'offline': { text: '📴 Offline', color: 'var(--fg2)', title: 'No connection to server' },
    'checking': { text: '⏳ Checking...', color: 'var(--fg2)', title: 'Checking sync status' }
  };

  const state = states[status] || states['checking'];
  el.textContent = state.text;
  el.style.color = state.color;
  el.title = state.title;

  // Add details to tooltip if data provided
  if (data && status === 'synced') {
    const details = [];
    if (data.pushed > 0) details.push(`${data.pushed} pushed`);
    if (data.pulled > 0) details.push(`${data.pulled} pulled`);
    if (data.conflicts > 0) details.push(`${data.conflicts} conflicts`);
    if (details.length > 0) {
      el.title = `Synced: ${details.join(', ')}`;
    }
  }
}

// ── Boot ─────────────────────────────────────────────────────────────────────
async function boot() {
  cfg = await fetch('/api/config').then(r=>r.json());
  document.getElementById('db-badge').textContent = 'DB: ' + cfg.db_type;
  updateDbSize();
  populateSelects();
  const saved = localStorage.getItem('wl-theme') || 'dark';
  themeIdx = themes.indexOf(saved); if (themeIdx<0) themeIdx=0;
  applyTheme(themes[themeIdx]);

  // Always show local data first (offline-first: instant load)
  await refreshList();
  document.addEventListener('paste', globalPaste);

  // v1.0.5: Add cleanup handler for page unload
  window.addEventListener('beforeunload', () => {
    stopAutoSave();
  });

  // Check if we need an initial cloud pull
  try {
    const init = await fetch('/api/sync/init').then(r=>r.json());
    if (init.online && !init.has_local_data) {
      toast('☁️ First launch — loading data from cloud…');
      const res  = await fetch('/api/sync/now', {method:'POST'});
      const data = await res.json();
      if (data.pulled > 0) {
        await refreshList();
        toast(`✅ Loaded ${data.pulled} record(s) from cloud`);
      }
    } else if (!init.has_local_data && !init.online) {
      toast('📴 Offline — no local data yet. Connect to sync from cloud.');
    }
  } catch(e) {}
}

function populateSelects() {
  ['f-status','f-status-field'].forEach(id=>{
    const sel = document.getElementById(id);
    if (!sel) return;
    sel.innerHTML = (id==='f-status'?'<option value="">All</option>':'') +
      cfg.status_options.map(s=>`<option>${s}</option>`).join('');
  });
  ['f-cat','f-category'].forEach(id=>{
    const sel = document.getElementById(id);
    if (!sel) return;
    sel.innerHTML = (id==='f-cat'?'<option value="">All</option>':'') +
      cfg.category_options.map(s=>`<option>${s}</option>`).join('');
  });
}

// ── Theme ─────────────────────────────────────────────────────────────────────
function cycleTheme() {
  themeIdx = (themeIdx+1) % themes.length;
  applyTheme(themes[themeIdx]);
  localStorage.setItem('wl-theme', themes[themeIdx]);
}
function applyTheme(mode) {
  let effective = mode;
  if (mode==='system')
    effective = window.matchMedia('(prefers-color-scheme:dark)').matches ? 'dark':'light';
  document.body.dataset.theme = effective;
  document.getElementById('theme-icon').innerHTML =
    mode==='dark'?'<i class="icon">🌙</i>': mode==='light'?'<i class="icon">☀️</i>':'<i class="icon">⚙️</i>';
  document.querySelector('meta[name=theme-color]').content =
    effective==='dark'?'#0f172a':'#f8fafc';
}

// ── Data ──────────────────────────────────────────────────────────────────────
async function refreshList() {
  const params = new URLSearchParams({
    customer:     document.getElementById('f-cust').value,
    project_name: document.getElementById('f-proj').value,
    status:       document.getElementById('f-status').value,
    category:     document.getElementById('f-cat').value,
    archive:      document.getElementById('f-arch').value,
  });
  records = await fetch('/api/records?'+params).then(r=>r.json());
  sortRecords(); renderTable();
}

function clearFilters() {
  ['f-cust','f-proj'].forEach(id=>document.getElementById(id).value='');
  ['f-status','f-cat','f-arch'].forEach(id=>document.getElementById(id).value='');
  refreshList();
}

function sortBy(col) {
  if (sortCol===col) sortAsc=!sortAsc;
  else { sortCol=col; sortAsc=true; }
  sortRecords(); renderTable();
}
function sortRecords() {
  records.sort((a,b)=>{
    const va=String(a[sortCol]||''), vb=String(b[sortCol]||'');
    return sortAsc ? va.localeCompare(vb) : vb.localeCompare(va);
  });
}

// ── Render ────────────────────────────────────────────────────────────────────
function statusBadge(s) {
  const cls = s==='Task Done'?'done': s==='WIP'||s==='Production'?'wip':
    s==='Cancelled'||s==='No More Activity'?'cancel':'other';
  return `<span class="badge badge-${cls}">${s||'–'}</span>`;
}
function esc(s) { return String(s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;'); }
function wlSnippet(wl) {
  if (!wl) return '';
  const clean = wl.replace(/\[IMG_B64:[^\]]+\]/g,'[img]').replace(/── .+? ──/g,'');
  return esc(clean.trim().slice(0,80)) + (clean.length>80?'…':'');
}

function renderTable() {
  const tbody = document.getElementById('records-body');
  tbody.innerHTML = records.map(r => {
    const archived = r.archive === 'Yes';
    const checked  = checkedIds.has(r.id);
    return `<tr class="${archived?'archived':''} ${r.id===selectedId?'selected':''}"
               onclick="selectRow(${r.id})" ondblclick="openEdit()">
      <td onclick="event.stopPropagation()" style="text-align:center">
        <input type="checkbox" ${checked?'checked':''} style="width:auto;cursor:pointer"
               onchange="toggleCheck(${r.id}, this.checked)">
      </td>
      <td>${r.week||''}</td>
      <td>${esc(r.category||'')}</td>
      <td>${esc(r.customer||'')}</td>
      <td>${esc(r.project_name||'')}</td>
      <td>${statusBadge(r.status)}</td>
      <td>${esc((r.task_summary||'').slice(0,60))}</td>
      <td class="wl-cell">${wlSnippet(r.worklogs)}</td>
      <td>${esc((r.project_schedule||'').slice(0,60))}</td>
      <td>${esc((r.todo_content||'').slice(0,40))}${r.todo_due_date?' 📅'+r.todo_due_date:''}</td>
      <td>${esc(r.application||'')}</td>
      <td>${esc(r.sso_modeln||'')}</td>
      <td class="mchp-device-cell">${esc(r.mchp_device||'')}</td>
      <td>${esc(r.ear||'')}</td>
      <td>${esc(r.bu||'')}</td>
      <td>${r.last_update||''}</td>
      <td>${r.create_date||''}</td>
      <td>${r.due_date||''}</td>
      <td>${r.archive||'No'}</td>
    </tr>`;
  }).join('');
  // Sync select-all checkbox state
  const allChk = document.getElementById('chk-all');
  if (allChk) {
    allChk.checked       = records.length > 0 && checkedIds.size === records.length;
    allChk.indeterminate = checkedIds.size > 0 && checkedIds.size < records.length;
  }
}

function selectRow(id) { selectedId = id; renderTable(); }

function toggleCheck(id, checked) {
  if (checked) checkedIds.add(id); else checkedIds.delete(id);
  renderTable();
}

function toggleAll(checked) {
  if (checked) records.forEach(r => checkedIds.add(r.id));
  else checkedIds.clear();
  renderTable();
}

// ── Modal helpers ─────────────────────────────────────────────────────────────
function openModal(title) {
  document.getElementById('modal-title').textContent = title;
  document.getElementById('modal-overlay').classList.add('open');
}
function closeModal() {
  // Check for unsaved changes
  if (hasUnsavedChanges()) {
    if (!confirm('You have unsaved changes. Close without saving?')) {
      return; // User canceled close
    }
  }

  // v1.0.5: Clear auto-save state (both interval and timer)
  stopAutoSave();
  initialFormState = null;

  document.getElementById('modal-overlay').classList.remove('open');
  editingId=null; pendingImages=[]; pendingFiles=[];
  document.getElementById('img-preview-area').innerHTML='';
}

// ── Auto-save State (v1.0.5: 30s periodic save with immediate sync) ──────────
let autoSaveTimer = null;
let autoSaveInterval = null;
let autoSaveInProgress = false;
let saveInProgress = false; // v1.0.5: Prevent race between auto-save and manual save
let initialFormState = null;
let syncDebounceTimer = null; // v1.0.5: Debounce sync calls

function captureFormState() {
  if (!editingId) return null; // Only track state for existing records
  return {
    task_summary: document.getElementById('f-task_summary')?.value || '',
    worklogs: document.getElementById('f-worklogs')?.value || '',
    project_schedule: document.getElementById('f-project_schedule')?.value || '',
    todo_content: document.getElementById('f-todo_content')?.value || '',
    archive: document.getElementById('f-archive')?.checked || false
  };
}

function hasUnsavedChanges() {
  if (!editingId || !initialFormState) return false;
  const current = captureFormState();
  return JSON.stringify(current) !== JSON.stringify(initialFormState);
}

function startAutoSave() {
  // v1.0.5: Start periodic auto-save every 30 seconds during editing
  if (!editingId) return;
  stopAutoSave(); // Clear any existing timer

  autoSaveInterval = setInterval(async () => {
    // v1.0.5: Safety checks to prevent race conditions and memory leaks
    if (!editingId || autoSaveInProgress || saveInProgress) return;
    if (!document.getElementById('modal-overlay')?.classList.contains('open')) {
      stopAutoSave(); // Cleanup if modal closed unexpectedly
      return;
    }
    if (!hasUnsavedChanges()) return; // Skip if no changes

    autoSaveInProgress = true;
    try {
      const data = formData();
      const res = await fetch(`/api/records/${editingId}`, {
        method: 'PUT',
        headers: {'Content-Type': 'application/json'},
        body: JSON.stringify(data)
      });

      if (res.ok) {
        initialFormState = captureFormState();
        const indicator = document.getElementById('auto-save-indicator');
        if (indicator) {
          indicator.textContent = '✓ Auto-saved at ' + new Date().toLocaleTimeString();
          indicator.style.opacity = '1';
          setTimeout(() => { indicator.style.opacity = '0'; }, 2000);
        }
        // v1.0.5: Trigger debounced sync after auto-save
        debouncedSync();
      }
    } catch (e) {
      console.error('Auto-save failed:', e);
    } finally {
      autoSaveInProgress = false;
    }
  }, 30000); // 30 seconds periodic auto-save
}

function stopAutoSave() {
  if (autoSaveInterval) {
    clearInterval(autoSaveInterval);
    autoSaveInterval = null;
  }
  if (autoSaveTimer) {
    clearTimeout(autoSaveTimer);
    autoSaveTimer = null;
  }
  if (syncDebounceTimer) {
    clearTimeout(syncDebounceTimer);
    syncDebounceTimer = null;
  }
  autoSaveInProgress = false;
  saveInProgress = false;
}

// v1.0.5: Debounced sync to prevent rapid-fire requests
function debouncedSync() {
  if (syncDebounceTimer) clearTimeout(syncDebounceTimer);
  syncDebounceTimer = setTimeout(() => {
    triggerSyncWithRetry();
  }, 1000); // Wait 1s before syncing
}

// v1.0.5: Sync with retry logic and better error handling
async function triggerSyncWithRetry(retries = 1) {
  updateSyncStatusBadge('syncing'); // Show syncing status
  try {
    const res = await fetch('/api/sync/now', {method: 'POST'});
    if (!res.ok) throw new Error(`Sync failed: ${res.status}`);
    const data = await res.json();

    // Update sync status indicator
    updateSyncStatusBadge('synced', data);
    return data;
  } catch (e) {
    console.error('Sync error:', e);

    if (retries > 0) {
      console.log(`Retrying sync... (${retries} attempts left)`);
      await new Promise(r => setTimeout(r, 2000)); // Wait 2s before retry
      return triggerSyncWithRetry(retries - 1);
    }

    // Show non-intrusive error to user after all retries failed
    updateSyncStatusBadge('offline');
    const indicator = document.getElementById('auto-save-indicator');
    if (indicator) {
      indicator.textContent = '⚠️ Sync pending (offline?)';
      indicator.style.opacity = '1';
      indicator.style.color = 'var(--orange)';
      setTimeout(() => {
        indicator.style.opacity = '0';
        indicator.style.color = '';
      }, 3000);
    }
  }
}

// Legacy function - now just starts the periodic auto-save
function scheduleAutoSave() {
  // Keep for compatibility, but now starts interval if not already running
  if (!autoSaveInterval && editingId) {
    startAutoSave();
  }
}

async function saveArchiveNow() {
  if (!editingId || autoSaveInProgress || saveInProgress) return;
  autoSaveInProgress = true;
  try {
    const data = formData();
    const res = await fetch(`/api/records/${editingId}`, {
      method: 'PUT',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify(data)
    });
    if (res.ok) {
      initialFormState = captureFormState();
      const indicator = document.getElementById('auto-save-indicator');
      if (indicator) {
        indicator.textContent = '✓ Saved';
        indicator.style.opacity = '1';
        setTimeout(() => { indicator.style.opacity = '0'; }, 1500);
      }
      const row = records.find(r => r.id === editingId);
      if (row) { row.archive = data.archive; renderTable(); }
      // v1.0.5: Trigger debounced sync after archive change
      debouncedSync();
    }
  } catch (e) {
    console.error('Archive save failed:', e);
  } finally {
    autoSaveInProgress = false;
  }
}

function dateFmt(s) {
  // Convert mm-dd-yyyy → yyyy-mm-dd for <input type=date>
  if (!s) return '';
  const m=s.match(/^(\d{2})-(\d{2})-(\d{4})$/);
  return m ? `${m[3]}-${m[1]}-${m[2]}` : s;
}
function dateOut(s) {
  // Convert yyyy-mm-dd → mm-dd-yyyy for API
  if (!s) return '';
  const m=s.match(/^(\d{4})-(\d{2})-(\d{2})$/);
  return m ? `${m[2]}-${m[3]}-${m[1]}` : s;
}

function autoWeek() {
  const d = document.getElementById('f-last_update').value;
  if (!d) return;
  const dt = new Date(d);
  const jan1 = new Date(dt.getFullYear(),0,1);
  const week = Math.ceil(((dt-jan1)/86400000 + jan1.getDay()+1)/7);
  document.getElementById('f-week').value = week;
}

function fillForm(rec) {
  const fd = (field, val) => { const el=document.getElementById('f-'+field); if(el) el.value=val||''; };
  fd('create_date', dateFmt(rec.create_date));
  fd('due_date',    dateFmt(rec.due_date));
  fd('last_update', dateFmt(rec.last_update));
  fd('customer',    rec.customer);
  fd('project_name',rec.project_name);
  fd('sso_modeln',  rec.sso_modeln);
  fd('ear',         rec.ear);
  fd('application', rec.application);
  fd('bu',          rec.bu);
  fd('mchp_device', rec.mchp_device);
  fd('week',        rec.week);
  fd('category',    rec.category||'General');
  fd('status-field',rec.status||'Not Started');
  fd('task_summary',rec.task_summary);
  fd('project_schedule',rec.project_schedule);
  fd('todo_content', rec.todo_content);
  fd('todo_due_date', dateFmt(rec.todo_due_date));
  document.getElementById('f-archive').checked = rec.archive==='Yes';
  // Character hints
  updateCharHint(document.getElementById('f-customer'),'ch-cust');
  updateCharHint(document.getElementById('f-project_name'),'ch-proj');
  updateCharHint(document.getElementById('f-sso_modeln'),'ch-sso');
  updateCharHint(document.getElementById('f-ear'),'ch-ear');
  updateCharHint(document.getElementById('f-application'),'ch-app');
  // Restore embedded images and files as UID placeholders + chips with delete buttons
  const area = document.getElementById('img-preview-area');
  area.innerHTML = '';
  pendingImages = [];
  pendingFiles = [];
  const imgPat  = /\[IMG_B64:([A-Za-z0-9+/=]+)\]/g;
  const filePat = /\[FILE:([^\]:\n]+):([^\]:\n]+):([A-Za-z0-9+/=]+)\]/g;
  let wlDisplay = rec.worklogs || '';
  let m;
  while ((m = imgPat.exec(rec.worklogs || '')) !== null) {
    const uid = 'img_' + Math.random().toString(36).slice(2, 10);
    pendingImages.push({uid, b64: m[1]});
    wlDisplay = wlDisplay.replace(m[0], `[📷 ${uid}]`);
    _addImageThumb(uid, m[1]);
  }
  while ((m = filePat.exec(rec.worklogs || '')) !== null) {
    const [, name, mimeType, b64] = m;
    const uid = 'file_' + Math.random().toString(36).slice(2, 10);
    const size = Math.round(b64.length * 0.75);
    pendingFiles.push({uid, name, mimeType, b64, size});
    wlDisplay = wlDisplay.replace(m[0], `[📎 ${uid}]`);
    _addFileChip(uid, name, size);
  }
  document.getElementById('f-worklogs').value = wlDisplay;

  // Capture initial state for change detection
  initialFormState = captureFormState();

  // v1.0.5: Start periodic auto-save every 30 seconds during editing
  if (editingId) {
    startAutoSave();
  }

  // Attach auto-save listeners to monitored fields (for immediate change detection)
  const autoSaveFields = ['f-task_summary', 'f-worklogs', 'f-project_schedule', 'f-todo_content'];
  autoSaveFields.forEach(fieldId => {
    const el = document.getElementById(fieldId);
    if (el) {
      el.removeEventListener('input', scheduleAutoSave); // Prevent duplicate listeners
      el.addEventListener('input', scheduleAutoSave);
    }
  });

  // Archive checkbox: save immediately (not debounced) to ensure DB is updated
  const archiveCheckbox = document.getElementById('f-archive');
  if (archiveCheckbox) {
    archiveCheckbox.removeEventListener('change', saveArchiveNow);
    archiveCheckbox.addEventListener('change', saveArchiveNow);
  }
}

function formData() {
  const gv = id => document.getElementById(id)?.value||'';
  // Replace pending placeholders with stored tokens for DB storage
  let wl = gv('f-worklogs');
  pendingImages.forEach(img => {
    wl = wl.split(`[📷 ${img.uid}]`).join(`[IMG_B64:${img.b64}]`);
  });
  pendingFiles.forEach(f => {
    const safeName = f.name.replace(/:/g, '_');
    wl = wl.split(`[📎 ${f.uid}]`).join(`[FILE:${safeName}:${f.mimeType}:${f.b64}]`);
  });
  return {
    create_date:  dateOut(gv('f-create_date')),
    due_date:     dateOut(gv('f-due_date')),
    last_update:  dateOut(gv('f-last_update')) || today(),
    week:         gv('f-week'),
    customer:     gv('f-customer').slice(0,20),
    project_name: gv('f-project_name').slice(0,20),
    sso_modeln:   gv('f-sso_modeln').slice(0,10),
    ear:          gv('f-ear').slice(0,20),
    application:  gv('f-application').slice(0,50),
    bu:           gv('f-bu'),
    mchp_device:  gv('f-mchp_device'),
    category:     gv('f-category'),
    status:       gv('f-status-field'),
    task_summary: gv('f-task_summary'),
    project_schedule: gv('f-project_schedule'),
    todo_content: gv('f-todo_content'),
    todo_due_date: dateOut(gv('f-todo_due_date')),
    worklogs:     wl,
    archive:      document.getElementById('f-archive').checked ? 'Yes':'No',
  };
}

function today() {
  const d=new Date(), p=n=>String(n).padStart(2,'0');
  return `${p(d.getMonth()+1)}-${p(d.getDate())}-${d.getFullYear()}`;
}

function openNew() {
  editingId=null; pendingImages=[]; pendingFiles=[];
  document.getElementById('img-preview-area').innerHTML='';
  fillForm({create_date:today(), last_update:today(), status:'Not Started', category:'General', archive:'No'});
  autoWeek();
  setWlTab('edit'); setPsTab('edit'); setTodoTab('edit');
  openModal('New Record');
}
function openEdit() {
  if (!selectedId) { toast('Select a record first'); return; }
  const rec = records.find(r=>r.id===selectedId);
  if (!rec) return;
  editingId = selectedId;
  fillForm(rec);
  setWlTab('edit'); setPsTab('edit'); setTodoTab('edit');
  openModal('Edit Record');
}

async function saveRecord() {
  // v1.0.5: Wait for any in-progress auto-save to complete
  if (autoSaveInProgress) {
    await new Promise(resolve => {
      const checkInterval = setInterval(() => {
        if (!autoSaveInProgress) {
          clearInterval(checkInterval);
          resolve();
        }
      }, 100);
      // Timeout after 5 seconds
      setTimeout(() => {
        clearInterval(checkInterval);
        resolve();
      }, 5000);
    });
  }

  saveInProgress = true; // Prevent auto-save during manual save
  const data = formData();
  const method = editingId ? 'PUT' : 'POST';
  const url    = editingId ? `/api/records/${editingId}` : '/api/records';
  try {
    const res = await fetch(url, {
      method,
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify(data)
    });
    const json = await res.json();
    if (!res.ok) {
      toast('❌ Save failed: ' + (json.error || res.status));
      console.error('Save error:', json);
      return;
    }

    // v1.0.5: Trigger immediate sync with retry after save
    const syncResult = await triggerSyncWithRetry(2);
    if (syncResult) {
      toast(editingId ? '✅ Record updated & synced' : '✅ Record added & synced');
    } else {
      toast(editingId ? '✅ Record updated (sync pending)' : '✅ Record added (sync pending)');
    }

    // Clear auto-save state and update initial state
    stopAutoSave();
    initialFormState = null; // Reset so closeModal doesn't warn
    closeModal();
    await refreshList();
  } catch (e) {
    toast('❌ Network error: ' + e.message);
    console.error(e);
  } finally {
    saveInProgress = false;
  }
}

async function deleteRecord() {
  if (!selectedId) { toast('Select a record first'); return; }
  if (!confirm('Delete this record?')) return;
  try {
    const res = await fetch(`/api/records/${selectedId}`, {method:'DELETE'});
    if (!res.ok) { toast('❌ Delete failed'); return; }
    selectedId = null;
    await refreshList();
    toast('🗑 Deleted');
  } catch(e) {
    toast('❌ Network error: ' + e.message);
  }
}

// ── Worklogs ──────────────────────────────────────────────────────────────────
function setWlTab(tab) {
  document.getElementById('wl-edit').style.display    = tab==='edit'?'block':'none';
  document.getElementById('wl-preview').style.display = tab==='preview'?'block':'none';
  if (tab==='preview') renderPreview();
  _syncTabBtns('wl-tabs', tab);
}
function renderPreview() {
  const md = document.getElementById('f-worklogs').value;
  document.getElementById('wl-preview').innerHTML = parseMarkdown(md);
}
function setPsTab(tab) {
  document.getElementById('ps-edit').style.display    = tab==='edit'?'block':'none';
  document.getElementById('ps-preview').style.display = tab==='preview'?'block':'none';
  if (tab==='preview') document.getElementById('ps-preview').innerHTML =
    parseMarkdown(document.getElementById('f-project_schedule').value);
  _syncTabBtns('ps-tabs', tab);
}
function setTodoTab(tab) {
  document.getElementById('todo-edit').style.display    = tab==='edit'?'block':'none';
  document.getElementById('todo-preview').style.display = tab==='preview'?'block':'none';
  if (tab==='preview') document.getElementById('todo-preview').innerHTML =
    parseMarkdown(document.getElementById('f-todo_content').value);
  _syncTabBtns('todo-tabs', tab);
}
function _syncTabBtns(containerId, activeTab) {
  const btns = document.getElementById(containerId)?.querySelectorAll('button');
  if (!btns) return;
  btns.forEach(b => {
    const isEdit = b.textContent.includes('Edit');
    const active = (activeTab==='edit') === isEdit;
    b.className = active ? 'btn btn-accent btn-sm' : 'btn btn-ghost btn-sm';
  });
}

function insertStamp() {
  // Show date-picker modal with today pre-filled
  const today = new Date();
  const pad = n => String(n).padStart(2,'0');
  document.getElementById('stamp-date').value =
    `${today.getFullYear()}-${pad(today.getMonth()+1)}-${pad(today.getDate())}`;
  document.getElementById('stamp-overlay').style.display = 'flex';
}
function closeStamp() {
  document.getElementById('stamp-overlay').style.display = 'none';
}
function confirmStamp() {
  const dateVal = document.getElementById('stamp-date').value;
  if (!dateVal) { closeStamp(); return; }
  const d   = new Date(dateVal + 'T00:00:00');
  const pad = n => String(n).padStart(2,'0');
  const stamp = `${d.getFullYear()}-${pad(d.getMonth()+1)}-${pad(d.getDate())}`;
  const ta = document.getElementById('f-worklogs');
  _insertAtCursor(ta, `\n── ${stamp} ──\n`);
  closeStamp();
  sortWorklogs();
  setWlTab('edit');
  ta.focus();
}

function sortWorklogs() {
  /** Parse all stamped entries, sort by date, reassemble. Pre-stamp text preserved. */
  const ta   = document.getElementById('f-worklogs');
  const text = ta.value;
  // Match both date-only and legacy date+time stamps
  const splitPat = /(── \d{4}-\d{2}-\d{2}(?:(?: \d{2}:\d{2}))? ──)/;
  const keyPat   = /── (\d{4}-\d{2}-\d{2})(?:(?: \d{2}:\d{2}))? ──/;
  const parts  = text.split(splitPat);
  const pre    = parts[0];
  const entries = [];
  for (let i = 1; i + 1 < parts.length; i += 2) {
    const stamp   = parts[i];
    const body    = parts[i+1] || '';
    const m       = stamp.match(keyPat);
    const dateKey = m ? m[1] : '';
    // Normalise stamp to date-only format
    const normStamp = dateKey ? `── ${dateKey} ──` : stamp;
    entries.push({ stamp: normStamp, body, dateKey });
  }
  if (entries.length < 2) return;
  entries.sort((a, b) => a.dateKey.localeCompare(b.dateKey));
  ta.value = pre + entries.map(e => e.stamp + e.body).join('');
}

function parseMarkdown(md) {
  if (!md) return '';

  // ── Pre-process: convert image/file tokens to sentinels ──
  // We do this before splitting into lines so multi-char tokens don't confuse things.
  const IMG_TOK = [];
  md = md.replace(/\[📷 (img_[a-z0-9]+)\]/g, (_, uid) => {
    const obj = pendingImages.find(p => p.uid === uid);
    IMG_TOK.push(obj ? `<img src="data:image/png;base64,${obj.b64}" class="img-thumb" onclick="openLightbox(this.src)" style="cursor:pointer">` : '<span style="color:var(--fg2)">[📷]</span>');
    return `\x00IMG${IMG_TOK.length - 1}\x00`;
  });
  md = md.replace(/\[IMG_B64:([A-Za-z0-9+/=]+)\]/g, (_, b64) => {
    IMG_TOK.push(`<img src="data:image/png;base64,${b64}" class="img-thumb" onclick="openLightbox(this.src)" style="cursor:pointer">`);
    return `\x00IMG${IMG_TOK.length - 1}\x00`;
  });

  const FILE_TOK = [];
  md = md.replace(/\[📎 (file_[a-z0-9]+)\]/g, (_, uid) => {
    const obj = pendingFiles.find(p => p.uid === uid);
    if (obj) {
      FILE_TOK.push(`<a href="data:${obj.mimeType};base64,${obj.b64}" download="${esc(obj.name)}" class="file-attach-link">📎 ${esc(obj.name)}</a>`);
    } else {
      FILE_TOK.push('<span style="color:var(--fg2)">[📎]</span>');
    }
    return `\x00FILE${FILE_TOK.length - 1}\x00`;
  });
  md = md.replace(/\[FILE:([^\]:\n]+):([^\]:\n]+):([A-Za-z0-9+/=]+)\]/g, (_, name, mime, b64) => {
    FILE_TOK.push(`<a href="data:${mime};base64,${b64}" download="${esc(name)}" class="file-attach-link">📎 ${esc(name)}</a>`);
    return `\x00FILE${FILE_TOK.length - 1}\x00`;
  });

  const lines = md.split('\n');

  // ── Build token list ───────────────────────────────────────────────────────
  // Each token: {type, indent, content, ordered}
  const tokens = [];
  for (const raw of lines) {
    const line = raw.replace(/\[IMG:[a-f0-9]{12}\]/g, '');

    if (/^── \d{4}-\d{2}-\d{2}(?: \d{2}:\d{2})? ──$/.test(line.trim())) {
      const norm = line.trim().replace(/── (\d{4}-\d{2}-\d{2})(?:(?: \d{2}:\d{2}))? ──/, '── $1 ──');
      tokens.push({type:'stamp', content: norm}); continue;
    }
    // Measure indent for every line (used by lists AND images)
    const indentMatch = line.match(/^([ \t]*)/);
    const rawIndent   = indentMatch ? indentMatch[1] : '';
    const indent      = rawIndent.replace(/\t/g, '  ').length;
    const rest        = line.slice(rawIndent.length);

    if (/^\x00IMG\d+\x00$/.test(rest.trim())) {
      tokens.push({type:'img', indent, content: rest.trim()}); continue;
    }
    if (/^\x00FILE\d+\x00$/.test(rest.trim())) {
      tokens.push({type:'file', indent, content: rest.trim()}); continue;
    }
    if (/^-{3,}$/.test(rest.trim())) {
      tokens.push({type:'hr'}); continue;
    }
    const hm = rest.match(/^(#{1,3})\s+(.*)/);
    if (hm) { tokens.push({type:'h', level: hm[1].length, content: hm[2]}); continue; }

    const bm = rest.match(/^[-*+]\s+(.*)/);
    if (bm) { tokens.push({type:'li', indent, ordered:false, content: bm[1]}); continue; }
    const nm = rest.match(/^(\d+)\.\s+(.*)/);
    if (nm) { tokens.push({type:'li', indent, ordered:true,  content: nm[2]}); continue; }

    tokens.push({type:'p', content: line});
  }

  // ── Render tokens → HTML ───────────────────────────────────────────────────
  // Stack entries: { tag:'ul'|'ol', level:number, liOpen:boolean }
  // liOpen tracks whether the current <li> is still open (sub-list may follow).
  // Sub-lists are emitted INSIDE the parent <li> (not as siblings), producing
  // valid HTML: <li>text<ul>...</ul></li>
  let html  = '';
  const stack = [];

  function peekTop() { return stack[stack.length - 1]; }

  function closeTop() {
    const top = stack[stack.length - 1];
    if (top.liOpen) { html += '</li>'; top.liOpen = false; }
    html += `</${stack.pop().tag}>`;
  }

  function closeAll()  { while (stack.length) closeTop(); }

  function closeTo(targetLevel) {
    while (stack.length && peekTop().level > targetLevel) closeTop();
  }

  for (const tok of tokens) {
    if (tok.type === 'li') {
      const tag   = tok.ordered ? 'ol' : 'ul';
      const level = Math.floor(tok.indent / 2);
      const top   = peekTop();

      if (!top) {
        html += `<${tag}>`;
        stack.push({tag, level, liOpen: false});
      } else if (level > top.level) {
        // Going deeper — open sub-list inside the current open <li>
        html += `<${tag}>`;
        stack.push({tag, level, liOpen: false});
      } else if (level < top.level) {
        closeTo(level);
        const cur = peekTop();
        if (!cur || cur.level !== level) {
          html += `<${tag}>`;
          stack.push({tag, level, liOpen: false});
        } else if (cur.tag !== tag) {
          closeTop();
          html += `<${tag}>`;
          stack.push({tag, level, liOpen: false});
        } else {
          if (cur.liOpen) { html += '</li>'; cur.liOpen = false; }
        }
      } else {
        // Same level
        if (top.tag !== tag) {
          closeTop();
          html += `<${tag}>`;
          stack.push({tag, level, liOpen: false});
        } else {
          if (top.liOpen) { html += '</li>'; top.liOpen = false; }
        }
      }
      html += `<li>${inlineM(tok.content)}`;
      peekTop().liOpen = true;
    } else {
      closeAll();
      if      (tok.type === 'stamp') html += `<div class="stamp">${esc(tok.content)}</div>`;
      else if (tok.type === 'img') {
        const marginEm = (tok.indent || 0) * 0.6;   // ~0.6em per space, same visual as list indent
        const style    = marginEm > 0 ? ` style="margin-left:${marginEm.toFixed(1)}em"` : '';
        const imgHtml  = tok.content.replace(/\x00IMG(\d+)\x00/, (_, i) => IMG_TOK[+i] || '');
        html += `<div${style}>${imgHtml}</div>`;
      }
      else if (tok.type === 'file') {
        const fileHtml = tok.content.replace(/\x00FILE(\d+)\x00/, (_, i) => FILE_TOK[+i] || '');
        html += `<div style="margin:4px 0">${fileHtml}</div>`;
      }
      else if (tok.type === 'hr')    html += '<hr>';
      else if (tok.type === 'h')     html += `<h${tok.level}>${inlineM(tok.content)}</h${tok.level}>`;
      else {
        let resolved = tok.content.replace(/\x00IMG(\d+)\x00/g, (_, i) => IMG_TOK[+i] || '');
        resolved = resolved.replace(/\x00FILE(\d+)\x00/g, (_, i) => FILE_TOK[+i] || '');
        const trimmed  = resolved.trim();
        if (trimmed) html += `<p>${inlineM(trimmed)}</p>`;
      }
    }
  }
  closeAll();
  return html;
}
function inlineM(s) {
  return esc(s)
    .replace(/\*\*(.+?)\*\*/g,'<strong>$1</strong>')
    .replace(/\*(.+?)\*/g,'<em>$1</em>')
    .replace(/__(.+?)__/g,'<u>$1</u>')
    .replace(/~~(.+?)~~/g,'<del>$1</del>')
    .replace(/`(.+?)`/g,'<code>$1</code>');
}

// ── Image handling ────────────────────────────────────────────────────────────
function globalPaste(e) {
  const modal = document.getElementById('modal-overlay');
  if (!modal.classList.contains('open')) return;
  const items = e.clipboardData?.items||[];
  for (const item of items) {
    if (item.type.startsWith('image/')) {
      e.preventDefault();
      const blob = item.getAsFile();
      readAndEmbedImage(blob); return;
    }
  }
}
function handleDrop(e) {
  e.preventDefault(); e.target.classList.remove('drag-over');
  const file = e.dataTransfer.files[0];
  if (file && file.type.startsWith('image/')) readAndEmbedImage(file);
}
function handleFileSelect(e) {
  const file = e.target.files[0];
  if (file) readAndEmbedImage(file);
  e.target.value='';
}
function readAndEmbedImage(blob) {
  const reader = new FileReader();
  reader.onload = async ev => {
    let b64 = ev.target.result.split(',')[1];
    try {
      const res = await fetch('/api/image/resize', {method:'POST',
        headers:{'Content-Type':'application/json'}, body:JSON.stringify({data:b64})});
      if (res.ok) b64 = (await res.json()).data;
    } catch(e) {}
    const uid = 'img_' + Math.random().toString(36).slice(2, 10);
    pendingImages.push({uid, b64});
    _addImageThumb(uid, b64);
    _insertAtCursor(document.getElementById('f-worklogs'), `\n[📷 ${uid}]\n`);
    toast('📷 Image attached');
  };
  reader.readAsDataURL(blob);
}

function _insertAtCursor(ta, text) {
  const start = ta.selectionStart;
  const end   = ta.selectionEnd;
  const before = ta.value.slice(0, start);
  const after  = ta.value.slice(end);
  ta.value = before + text + after;
  // Move cursor to after the inserted text
  const newPos = start + text.length;
  ta.selectionStart = ta.selectionEnd = newPos;
  ta.focus();
}

function _addImageThumb(uid, b64) {
  const area = document.getElementById('img-preview-area');
  const thumb = document.createElement('div');
  thumb.id = 'thumb_' + uid;
  thumb.style.cssText = 'display:inline-flex;align-items:center;gap:6px;' +
    'background:var(--bg3);border-radius:6px;padding:4px 8px;margin:3px;font-size:11px;' +
    'position:relative;';
  const img = document.createElement('img');
  img.style.cssText = 'width:48px;height:36px;object-fit:cover;border-radius:4px;cursor:default;';
  img.src = 'data:image/png;base64,' + b64;
  const lbl = document.createElement('span');
  lbl.style.color = 'var(--fg2)'; lbl.textContent = '📷';
  const del = document.createElement('button');
  del.textContent = '✕';
  del.title = 'Remove image';
  del.style.cssText = 'border:none;background:var(--red);color:#fff;border-radius:4px;' +
    'width:18px;height:18px;cursor:pointer;font-size:10px;padding:0;line-height:1;' +
    'flex-shrink:0;';
  del.onclick = () => removeImage(uid);
  thumb.appendChild(img); thumb.appendChild(lbl); thumb.appendChild(del);
  area.appendChild(thumb);
}

function removeImage(uid) {
  // Remove from pendingImages
  pendingImages = pendingImages.filter(p => p.uid !== uid);
  // Remove thumbnail
  const t = document.getElementById('thumb_' + uid);
  if (t) t.remove();
  // Remove placeholder from textarea
  const ta = document.getElementById('f-worklogs');
  ta.value = ta.value.replace(new RegExp(`\\n?\\[📷 ${uid}\\]\\n?`, 'g'), '\n').trim();
  toast('Image removed');
}

// ── Non-image file handling ───────────────────────────────────────────────────
const FILE_SIZE_LIMIT = 5 * 1024 * 1024; // 5 MB

function handleFileDrop(e) {
  e.preventDefault(); e.target.classList.remove('drag-over');
  const file = e.dataTransfer.files[0];
  if (!file) return;
  if (file.type.startsWith('image/')) { readAndEmbedImage(file); return; }
  readAndEmbedFile(file);
}
function handleNonImageFileSelect(e) {
  const file = e.target.files[0];
  if (file) readAndEmbedFile(file);
  e.target.value = '';
}
function readAndEmbedFile(file) {
  if (file.size > FILE_SIZE_LIMIT) {
    toast('File too large (max 5 MB)'); return;
  }
  const reader = new FileReader();
  reader.onload = ev => {
    const b64 = ev.target.result.split(',')[1];
    const uid  = 'file_' + Math.random().toString(36).slice(2, 10);
    const mimeType = file.type || 'application/octet-stream';
    pendingFiles.push({uid, name: file.name, mimeType, b64, size: file.size});
    _addFileChip(uid, file.name, file.size);
    _insertAtCursor(document.getElementById('f-worklogs'), `\n[📎 ${uid}]\n`);
    toast('📎 File attached');
  };
  reader.readAsDataURL(file);
}
function _addFileChip(uid, name, size) {
  const area = document.getElementById('img-preview-area');
  const chip = document.createElement('div');
  chip.id = 'thumb_' + uid;
  chip.style.cssText = 'display:inline-flex;align-items:center;gap:6px;' +
    'background:var(--bg3);border-radius:6px;padding:4px 8px;margin:3px;font-size:11px;';
  const sizeKB = (size / 1024).toFixed(1);
  const lbl = document.createElement('span');
  lbl.style.color = 'var(--fg)';
  lbl.textContent = `📎 ${name} (${sizeKB} KB)`;
  const del = document.createElement('button');
  del.textContent = '✕';
  del.title = 'Remove file';
  del.style.cssText = 'border:none;background:var(--red);color:#fff;border-radius:4px;' +
    'width:18px;height:18px;cursor:pointer;font-size:10px;padding:0;line-height:1;flex-shrink:0;';
  del.onclick = () => removeFile(uid);
  chip.appendChild(lbl); chip.appendChild(del);
  area.appendChild(chip);
}
function removeFile(uid) {
  pendingFiles = pendingFiles.filter(f => f.uid !== uid);
  const t = document.getElementById('thumb_' + uid);
  if (t) t.remove();
  const ta = document.getElementById('f-worklogs');
  ta.value = ta.value.replace(new RegExp(`\\n?\\[📎 ${uid}\\]\\n?`, 'g'), '\n').trim();
  toast('File removed');
}

// ── Import ────────────────────────────────────────────────────────────────────
function openImport()  { document.getElementById('import-overlay').style.display='flex'; }
function closeImport() {
  document.getElementById('import-overlay').style.display='none';
  document.getElementById('import-result').innerHTML='';
  document.getElementById('import-db-file').value = '';
}

async function doImport() {
  const fileInput = document.getElementById('import-db-file');
  const file = fileInput.files[0];

  if (!file) { toast('請選擇 .db 檔案'); return; }

  const div = document.getElementById('import-result');
  div.innerHTML = '<div style="margin:8px 0;font-size:12px;color:var(--yellow)">⏳ 正在讀取並匯入資料...</div>';

  try {
    const formData = new FormData();
    formData.append('db_file', file);

    const res = await fetch('/api/import', {
      method: 'POST',
      body: formData
    });

    const result = await res.json();

    if (!res.ok) {
      div.innerHTML = `<div style="margin:8px 0;font-size:12px;color:var(--red)">❌ 匯入失敗: ${result.error || '未知錯誤'}</div>`;
      toast('❌ 匯入失敗');
      return;
    }

    div.innerHTML = `<div style="margin:8px 0;font-size:12px">
      ➕ 新增: <strong>${result.added}</strong> &nbsp;
      🔀 合併: <strong>${result.merged}</strong> &nbsp;
      ⏭ 略過: <strong>${result.skipped}</strong>
    </div>` + (result.results || []).map(r=>`
      <div class="ir-row ir-${r.status}">
        <span>${r.status==='added'?'➕':r.status==='merged'?'🔀':'⏭'}</span>
        <span>${esc(r.customer)}</span>
        <span style="color:var(--fg2)">${esc(r.project)}</span>
      </div>`).join('');

    toast(`✅ 匯入完成: ${result.added} 新增, ${result.merged} 合併`);
    refreshList();
  } catch (err) {
    div.innerHTML = `<div style="margin:8px 0;font-size:12px;color:var(--red)">❌ 處理錯誤: ${err.message}</div>`;
    toast('❌ 處理錯誤');
  }
}

// ── Word export ───────────────────────────────────────────────────────────────
let exportIds = [];

async function exportWord() {
  const ids = [...checkedIds];
  if (ids.length === 0) { toast('☑ Check at least one record first'); return; }
  exportIds = ids;
  await openWeekSelect();
}

// ── Excel Export Modal ───────────────────────────────────────────────────────
let excelExportIds = [];

async function exportExcel() {
  const ids = [...checkedIds];
  if (ids.length === 0) { toast('☑ Check at least one record first'); return; }
  excelExportIds = ids;
  openExcelSelect();
}

function openExcelSelect() {
  document.getElementById('excel-overlay').style.display = 'flex';

  // Setup format change listener to show/hide date filter
  const wlrRadio = document.getElementById('format-wlr');
  const todoRadio = document.getElementById('format-todo');
  const dateFilterSection = document.getElementById('date-filter-section');

  const updateDateFilter = () => {
    dateFilterSection.style.display = todoRadio.checked ? 'block' : 'none';
  };

  wlrRadio.addEventListener('change', updateDateFilter);
  todoRadio.addEventListener('change', updateDateFilter);
  updateDateFilter(); // Initial state

  // Visual feedback for selected format
  const wlrLabel = document.getElementById('format-wlr-label');
  const todoLabel = document.getElementById('format-todo-label');

  const updateLabels = () => {
    wlrLabel.style.borderColor = wlrRadio.checked ? 'var(--accent)' : 'transparent';
    todoLabel.style.borderColor = todoRadio.checked ? 'var(--accent)' : 'transparent';
  };

  wlrRadio.addEventListener('change', updateLabels);
  todoRadio.addEventListener('change', updateLabels);
  updateLabels();
}

function closeExcelSelect() {
  document.getElementById('excel-overlay').style.display = 'none';
  excelExportIds = [];
}

async function confirmExcelExport() {
  const format = document.querySelector('input[name="excel-format"]:checked').value;
  const dateFilter = document.querySelector('input[name="date-filter"]:checked').value;

  // Save excelExportIds before closing (which clears excelExportIds)
  const idsToExport = [...excelExportIds];
  closeExcelSelect();
  toast(`Generating Excel export...`);

  try {
    const res = await fetch('/api/export/excel', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({
        ids: idsToExport,
        format: format,
        date_filter: dateFilter
      })
    });

    if (!res.ok) {
      const err = await res.json().catch(() => ({error: res.statusText}));
      toast('❌ Export failed: ' + (err.error || res.status));
      return;
    }

    const blob = await res.blob();
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url;

    const formatPrefix = format === 'WLR' ? 'WorkLog_WLR' : 'WorkLog_ToDo';
    a.download = `${formatPrefix}_${new Date().toISOString().slice(0,10)}.xlsx`;
    a.click();
    URL.revokeObjectURL(url);
    toast('✅ Excel exported successfully');
  } catch (e) {
    toast('❌ ' + e.message);
  }
}

async function exportPdf() {
  const ids = [...checkedIds];
  if (ids.length === 0) { toast('☑ Check at least one record first'); return; }
  try {
    const res = await fetch('/api/export/pdf', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ ids })
    });
    if (!res.ok) throw new Error(await res.text());
    const blob = await res.blob();
    const a = document.createElement('a');
    a.href = URL.createObjectURL(blob);
    a.download = `WorkLog_${new Date().toISOString().slice(0,10)}.pdf`;
    a.click();
    toast('📑 PDF exported successfully');
  } catch (e) {
    toast('❌ PDF export failed: ' + e.message);
  }
}

async function openWeekSelect() {
  const sel = document.getElementById('week-select');
  const info = document.getElementById('week-info');
  sel.innerHTML = '<option value="loading">Loading weeks...</option>';
  info.style.display = 'none';
  document.getElementById('week-overlay').style.display = 'flex';

  try {
    const res = await fetch('/api/export/weeks', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ ids: exportIds })
    });
    const weeks = await res.json();

    sel.innerHTML = '<option value="all">📋 All Worklogs (no filter)</option>';
    weeks.forEach(w => {
      const opt = document.createElement('option');
      opt.value = w.value;
      opt.textContent = `${w.label}  (${w.range})`;
      sel.appendChild(opt);
    });

    if (weeks.length === 0) {
      info.textContent = 'No dated worklogs found in selected records.';
      info.style.display = 'block';
    }
  } catch (e) {
    sel.innerHTML = '<option value="all">All Worklogs</option>';
    info.textContent = 'Could not load weeks: ' + e.message;
    info.style.display = 'block';
  }
}

function closeWeekSelect() {
  document.getElementById('week-overlay').style.display = 'none';
  exportIds = [];
}

async function confirmWeekExport() {
  const week = document.getElementById('week-select').value;
  if (week === 'loading') return;

  // Save exportIds before closing (which clears exportIds)
  const idsToExport = [...exportIds];
  closeWeekSelect();
  toast(`Generating report for ${idsToExport.length} record(s)…`);

  try {
    const res = await fetch('/api/export/word', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ ids: idsToExport, week })
    });
    if (!res.ok) {
      const err = await res.json().catch(() => ({error: res.statusText}));
      toast('❌ Export failed: ' + (err.error || res.status));
      return;
    }
    const blob = await res.blob();
    const url  = URL.createObjectURL(blob);
    const a    = document.createElement('a');
    a.href = url;
    a.download = `WorkLog_${new Date().toISOString().slice(0,10)}.docx`;
    a.click();
    URL.revokeObjectURL(url);
    toast('✅ Report downloaded');
  } catch (e) {
    toast('❌ ' + e.message);
  }
}

// ── Misc ──────────────────────────────────────────────────────────────────────
function updateCharHint(el, hintId) {
  const h=document.getElementById(hintId); if(!h)return;
  h.textContent=`${el.value.length}/${el.maxLength}`;
  h.style.color=el.value.length>=el.maxLength?'var(--red)':'var(--fg2)';
}
let toastTimer;
function toast(msg) {
  const el=document.getElementById('toast'); el.textContent=msg;
  el.classList.add('show'); clearTimeout(toastTimer);
  toastTimer=setTimeout(()=>el.classList.remove('show'),2600);
}

// ── Sync & Offline ────────────────────────────────────────────────────────────
let _syncInterval = null;
let _dbSizeInterval = null;

async function updateDbSize() {
  try {
    const res = await fetch('/api/db/size');
    const data = await res.json();
    document.getElementById('db-size-badge').textContent = `DB: ${data.size}`;
  } catch (e) {
    document.getElementById('db-size-badge').textContent = 'DB: --';
  }
}

function startSyncPoller() {
  updateSyncStatus();
  _syncInterval = setInterval(updateSyncStatus, 15000);
  updateDbSize();
  _dbSizeInterval = setInterval(updateDbSize, 60000);
}

async function updateSyncStatus() {
  try {
    const res  = await fetch('/api/sync/status');
    const data = await res.json();
    const dot  = document.getElementById('sync-dot');
    const txt  = document.getElementById('sync-text');
    const pend = document.getElementById('sync-pending');
    const badgeWrap = document.getElementById('conflict-badge-wrap');

    if (data.online) {
      dot.className = 'sync-dot sync-online';
      txt.textContent = data.last_sync
        ? `Online — last sync ${data.last_sync}` : 'Online — syncing…';
      // v1.0.5: Update footer sync badge
      if (data.pending > 0) {
        updateSyncStatusBadge('pending');
      } else {
        updateSyncStatusBadge('synced');
      }
    } else {
      dot.className = 'sync-dot sync-offline';
      txt.textContent = 'Offline — changes saved locally';
      updateSyncStatusBadge('offline');
    }
    pend.textContent = data.pending > 0
      ? `  ${data.pending} pending upload(s)` : '';
    if (data.conflicts > 0) {
      badgeWrap.innerHTML =
        `<span class="conflict-badge" onclick="openConflicts()">
          ⚠ ${data.conflicts} conflict${data.conflicts>1?'s':''}
        </span>`;
    } else {
      badgeWrap.innerHTML = '';
    }
  } catch(e) {
    const dot = document.getElementById('sync-dot');
    if (dot) dot.className = 'sync-dot sync-error';
    updateSyncStatusBadge('offline');
  }
}

async function openSyncSettings() {
  const res = await fetch('/api/env');
  const d   = await res.json();
  const pgUrl = d.POSTGRES_URL || '';
  const pbUrl = d.POCKETBASE_URL || '';
  if (pgUrl) {
    document.getElementById('ss-mode').value = 'postgres';
    document.getElementById('ss-pg-url').value = pgUrl;
  } else if (pbUrl) {
    document.getElementById('ss-mode').value = 'pocketbase';
    document.getElementById('ss-pb-url').value = pbUrl;
    document.getElementById('ss-pb-token').value = '';
  } else {
    document.getElementById('ss-mode').value = 'local';
  }
  onSsModeChange();
  document.getElementById('sync-settings-overlay').style.display = 'flex';
}

function closeSyncSettings() {
  document.getElementById('sync-settings-overlay').style.display = 'none';
}

function onSsModeChange() {
  const mode = document.getElementById('ss-mode').value;
  document.getElementById('ss-pb-section').style.display = mode === 'pocketbase' ? '' : 'none';
  document.getElementById('ss-pg-section').style.display = mode === 'postgres'   ? '' : 'none';
}

async function saveSyncSettings() {
  const mode = document.getElementById('ss-mode').value;
  const payload = { POCKETBASE_URL: '', POCKETBASE_TOKEN: '', POSTGRES_URL: '' };
  if (mode === 'pocketbase') {
    payload.POCKETBASE_URL   = document.getElementById('ss-pb-url').value.trim();
    payload.POCKETBASE_TOKEN = document.getElementById('ss-pb-token').value.trim();
  } else if (mode === 'postgres') {
    payload.POSTGRES_URL = document.getElementById('ss-pg-url').value.trim();
  }
  const res = await fetch('/api/env', { method: 'POST', headers: {'Content-Type':'application/json'}, body: JSON.stringify(payload) });
  const d   = await res.json();
  if (d.ok) {
    closeSyncSettings();
    toast('✅ 設定已儲存，正在重新連線…');
    setTimeout(checkSyncStatus, 800);
  } else {
    toast('❌ 儲存失敗');
  }
}

async function syncNow() {
  const btn = document.getElementById('sync-btn');
  btn.textContent = '↻ Syncing…'; btn.disabled = true;
  try {
    // First test connectivity and show diagnostics if not configured
    const testRes  = await fetch('/api/sync/test');
    const testData = await testRes.json();
    const isPostgres = testData.backend === 'postgres';

    if (!isPostgres) {
      // PocketBase-specific checks
      if (!testData.has_pocketbase_lib) {
        toast('❌ pocketbase library missing — run: pip install pocketbase'); return;
      }
      if (!testData.configured) {
        toast('⚙️ No cloud database configured — set credentials in Settings'); return;
      }
    }
    if (!testData.configured) {
      toast('⚙️ No cloud database configured — set credentials in Settings'); return;
    }
    if (!testData.online) {
      const detail = testData.error ? `\n${testData.error}` : '';
      toast(`📴 Cannot reach ${testData.url} — check credentials or network${detail}`); return;
    }

    // Schema check only for PocketBase
    if (!isPostgres) {
      const schemaRes = await fetch('/api/sync/check-schema');
      const schemaData = await schemaRes.json();
      if (schemaData.ok === false && schemaData.missing_columns) {
        if (confirm(`⚠️ PocketBase database needs update!\n\nMissing columns: ${schemaData.missing_columns.join(', ')}\n\nContinue sync anyway? (May cause errors)`)) {
          // User chose to continue despite schema mismatch
        } else {
          return; // User cancelled
        }
      }
    }

    const res  = await fetch('/api/sync/now', {method:'POST'});
    const data = await res.json();
    if (data.ok === false) {
      toast('📴 Offline — changes queued');
    } else {
      const msg = [];
      if (data.pushed)    msg.push(`${data.pushed} uploaded`);
      if (data.pulled)    msg.push(`${data.pulled} downloaded`);
      if (data.conflicts) msg.push(`⚠ ${data.conflicts} conflict(s)`);
      if (data.errors && data.errors.length) {
        console.error('Sync errors:', data.errors);
        msg.push(`${data.errors.length} error(s)`);
      }
      toast(msg.length ? `✅ Sync: ${msg.join(', ')}` : '✅ Up to date');
      await refreshList();
      await updateSyncStatus();
      if (data.conflicts) openConflicts();
    }
  } catch(e) {
    const detail = e.name && e.name !== 'Error' ? ` [${e.name}]` : '';
    toast('❌ Sync error: ' + e.message + detail);
    console.error('syncNow error', e);
  }
  finally { btn.textContent = '↻ Sync'; btn.disabled = false; }
}

// ── Conflict Resolution ───────────────────────────────────────────────────────
async function openConflicts() {
  const res       = await fetch('/api/sync/conflicts');
  const conflicts = await res.json();
  const list      = document.getElementById('conflict-list');

  if (!conflicts.length) {
    list.innerHTML = '<p style="color:var(--fg2);font-size:13px">No conflicts — all clear ✅</p>';
    document.getElementById('conflict-overlay').classList.add('open');
    return;
  }

  const fields = ['customer','project_name','status','category',
                  'task_summary','work_details','last_update','due_date'];

  function fieldRow(label, lv, cv, fieldName, diffFields) {
    const diff = (lv||'') !== (cv||'');
    const isChanged = diffFields.includes(fieldName); // v1.0.5: check if field is in diff list
    const highlightStyle = isChanged ? 'border-left:3px solid #f59e0b;padding-left:8px;background:rgba(245,158,11,0.1)' : '';
    return `<div class="conf-field" style="${highlightStyle}"><strong>${label}:</strong>
      <span class="${diff?'conf-diff':''}">${esc(lv||'—')}</span>
      ${diff ? `<span style="color:var(--fg2)"> ↔ </span>
      <span class="conf-diff" style="background:rgba(56,189,248,0.15)">
        ${esc(cv||'—')}</span>` : ''}
    </div>`;
  }

  list.innerHTML = conflicts.map(c => {
    const cloud = c.cloud_snapshot || {};
    const isUnresolvable = c.conflict_type === 'unresolvable';
    // v1.0.5: Extract field diff metadata
    const diffFields = cloud._diff_fields || [];
    const diffCount = diffFields.length;
    const warningMsg = isUnresolvable
      ? `<div style="background:#7c2d2d;color:#fca5a5;padding:8px 12px;border-radius:6px;margin-bottom:10px;font-size:11px">
          ⚠️ Unable to determine which version is newer. Please choose manually or keep both as backup.
        </div>` : '';
    return `<div class="conflict-card">
      <div style="font-size:13px;font-weight:700;margin-bottom:10px">
        📋 ${esc(c.customer||'?')} · ${esc(c.project_name||'?')}
        <span style="font-size:10px;color:var(--fg2);font-weight:400;margin-left:8px">
          detected ${esc(c.detected_at||'')}
        </span>
        ${isUnresolvable ? '<span style="background:#ef4444;color:#fff;padding:2px 6px;border-radius:4px;font-size:9px;margin-left:6px">NEEDS DECISION</span>' : ''}
        ${diffCount > 0 ? `<span style="background:#f59e0b;color:#fff;padding:2px 6px;border-radius:4px;font-size:9px;margin-left:6px">${diffCount} field${diffCount>1?'s':''} changed</span>` : ''}
      </div>
      ${warningMsg}
      <div style="display:flex;gap:16px">
        <div class="conflict-col local-col">
          <h4>📝 My offline edits</h4>
          ${fields.map(f => fieldRow(f, c[f], cloud[f], f, diffFields)).join('')}
        </div>
        <div class="conflict-col cloud-col">
          <h4>☁️ Cloud version</h4>
          ${fields.map(f => fieldRow(f, cloud[f], c[f], f, diffFields)).join('')}
        </div>
      </div>
      <div style="display:flex;gap:8px;margin-top:12px">
        <button class="btn btn-yellow btn-sm"
                onclick="resolveConflict(${c.conflict_id},'local')">
          <i class="icon">✅</i> Keep My Version
        </button>
        <button class="btn btn-accent btn-sm"
                onclick="resolveConflict(${c.conflict_id},'cloud')">
          <i class="icon">☁️</i> Use Cloud Version
        </button>
        ${isUnresolvable ? `<button class="btn btn-ghost btn-sm"
                onclick="resolveConflict(${c.conflict_id},'backup')">
          <i class="icon">💾</i> Keep Both (Backup)
        </button>` : ''}
      </div>
    </div>`;
  }).join('');

  document.getElementById('conflict-overlay').classList.add('open');
}

function closeConflicts() {
  document.getElementById('conflict-overlay').classList.remove('open');
}

async function resolveConflict(conflictId, keep) {
  const res = await fetch(`/api/sync/conflicts/${conflictId}`, {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({keep})
  });
  if (res.ok) {
    toast(keep==='local' ? '✅ Kept local version' : '☁️ Using cloud version');
    await openConflicts();
    await refreshList();
    await updateSyncStatus();
  }
}

document.addEventListener('click', e => {
  if (e.target.closest('#conflict-overlay') === document.getElementById('conflict-overlay')
      && !e.target.closest('#conflict-modal')) closeConflicts();
});


document.addEventListener('keydown', e => {
  if (e.key === 'Escape') closeLightbox();
});

function openLightbox(src) {
  const lb = document.getElementById('lightbox');
  const container = document.getElementById('lightbox-container');
  const img = document.getElementById('lightbox-img');
  img.src = src;
  // Reset container size
  container.style.width = '600px';
  container.style.height = '450px';
  lb.classList.add('open');
}

function closeLightbox(e) {
  // If called from backdrop click, only close when clicking the backdrop itself
  if (e && e.target !== document.getElementById('lightbox') &&
      e.target !== document.getElementById('lightbox-close')) return;
  document.getElementById('lightbox').classList.remove('open');
  document.getElementById('lightbox-img').src = '';
}

// Mouse wheel zoom for lightbox image
document.addEventListener('DOMContentLoaded', function() {
  const container = document.getElementById('lightbox-container');
  if (container) {
    container.addEventListener('wheel', function(e) {
      e.preventDefault();
      const currentWidth = container.offsetWidth;
      const currentHeight = container.offsetHeight;
      const delta = e.deltaY > 0 ? 0.9 : 1.1; // Zoom out or in

      const newWidth = Math.max(200, Math.min(currentWidth * delta, window.innerWidth * 0.95));
      const newHeight = Math.max(150, Math.min(currentHeight * delta, window.innerHeight * 0.95));

      container.style.width = newWidth + 'px';
      container.style.height = newHeight + 'px';
    });
  }
});

boot();
startSyncPoller();

if ('serviceWorker' in navigator) {
  navigator.serviceWorker.register('/sw.js', { scope: '/' }).catch(() => {});
}
</script>
</body>
</html>"""

@app.route("/")
def index():
    html_with_version = HTML.replace('{{APP_VERSION}}', APP_VERSION)
    return render_template_string(html_with_version)

@app.route("/manifest.json")
def manifest():
    return jsonify({
        "name":             "Work Log Journal",
        "short_name":       "Work Log",
        "start_url":        "/",
        "display":          "standalone",
        "background_color": "#0f172a",
        "theme_color":      "#0f172a",
        "icons": [
            {"src": "/icon-192.png",       "sizes": "192x192", "type": "image/png"},
            {"src": "/icon-512.png",       "sizes": "512x512", "type": "image/png"},
            {"src": "/apple-touch-icon.png", "sizes": "180x180", "type": "image/png"},
        ]
    })


def _make_icon_png(size: int) -> bytes:
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new("RGBA", (size, size), (15, 23, 42, 255))   # #0f172a bg
    draw = ImageDraw.Draw(img)
    m = size // 8
    draw.rounded_rectangle(
        [m, m, size - m, size - m],
        radius=size // 5,
        fill=(56, 189, 248, 255),   # #38bdf8 accent
    )
    font_size = size // 3
    font = None
    for path in [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]:
        try:
            font = ImageFont.truetype(path, font_size)
            break
        except OSError:
            continue
    if font is None:
        font = ImageFont.load_default()
    text = "WL"
    bbox = draw.textbbox((0, 0), text, font=font)
    x = (size - (bbox[2] - bbox[0])) // 2 - bbox[0]
    y = (size - (bbox[3] - bbox[1])) // 2 - bbox[1]
    draw.text((x, y), text, fill=(15, 23, 42, 255), font=font)
    buf = io.BytesIO()
    img.save(buf, "PNG")
    buf.seek(0)
    return buf


@app.route("/icon-192.png")
def icon_192():
    return send_file(_make_icon_png(192), mimetype="image/png",
                     max_age=86400, as_attachment=False)


@app.route("/icon-512.png")
def icon_512():
    return send_file(_make_icon_png(512), mimetype="image/png",
                     max_age=86400, as_attachment=False)


@app.route("/apple-touch-icon.png")
@app.route("/apple-touch-icon-precomposed.png")
def apple_touch_icon():
    return send_file(_make_icon_png(180), mimetype="image/png",
                     max_age=86400, as_attachment=False)


_SW_JS = """\
const CACHE = 'worklog-v2';

self.addEventListener('install', e => {
  e.waitUntil(caches.open(CACHE).then(c => c.addAll(['/'])));
  self.skipWaiting();
});

self.addEventListener('activate', e => {
  e.waitUntil(
    caches.keys().then(keys =>
      Promise.all(keys.filter(k => k !== CACHE).map(k => caches.delete(k)))
    )
  );
  self.clients.claim();
});

self.addEventListener('fetch', e => {
  if (e.request.method !== 'GET') {
    e.respondWith(fetch(e.request));
    return;
  }
  if (e.request.url.includes('/api/')) {
    e.respondWith(fetch(e.request).catch(() => caches.match(e.request)));
    return;
  }
  e.respondWith(
    caches.match(e.request).then(cached => cached || fetch(e.request))
  );
});
"""


@app.route("/sw.js")
def service_worker():
    resp = make_response(_SW_JS)
    resp.headers["Content-Type"] = "application/javascript"
    resp.headers["Service-Worker-Allowed"] = "/"
    resp.headers["Cache-Control"] = "no-cache"
    return resp

if __name__ == "__main__":
    host = os.environ.get("HOST", "0.0.0.0")
    port = int(os.environ.get("PORT", 5000))

    # Startup diagnostics
    env_file = Path(__file__).parent / ".env"
    print(f"\n  Work Log Journal - http://localhost:{port}")
    print(f"  Local network:     http://<your-ip>:{port}")
    print(f"  .env file:         {'FOUND [OK] ' + str(env_file) if env_file.exists() else 'NOT FOUND - create .env next to app.py'}")
    print(f"  POCKETBASE_URL:      {POCKETBASE_URL[:40] + '...' if POCKETBASE_URL else '(not set)'}")
    print(f"  POCKETBASE_TOKEN:      {'set [OK]' if POCKETBASE_TOKEN else '(not set)'}")
    print(f"  pocketbase library:  {'installed [OK]' if HAS_POCKETBASE else 'MISSING - run: pip install pocketbase'}")
    if HAS_POCKETBASE and POCKETBASE_URL and POCKETBASE_TOKEN:
        print(f"  DB:                PocketBase (cloud) [OK]")
    else:
        reasons = []
        if not HAS_POCKETBASE:   reasons.append("pocketbase not installed")
        if not POCKETBASE_URL:   reasons.append("POCKETBASE_URL missing")
        if not POCKETBASE_TOKEN:   reasons.append("POCKETBASE_TOKEN missing")
        print(f"  DB:                SQLite (local) - {', '.join(reasons)}")
    print()
    app.run(host=host, port=port, debug=False)
